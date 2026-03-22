"""OCR and lightweight text extraction helpers for research corpora."""

from __future__ import annotations

import csv
import html
import json
import os
import re
import shutil
import sys
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List

try:
    import docx as python_docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    python_docx = None
    PYTHON_DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    pd = None
    PANDAS_AVAILABLE = False

try:
    import pythoncom
    import win32com.client
    WIN32COM_AVAILABLE = True
except ImportError:
    pythoncom = None
    win32com = None
    WIN32COM_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    TESSERACT_PYTHON_AVAILABLE = True
except ImportError:
    pytesseract = None
    Image = None
    TESSERACT_PYTHON_AVAILABLE = False

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    convert_from_path = None
    PDF2IMAGE_AVAILABLE = False


def _ai_agents_root() -> Path:
    return Path(__file__).resolve().parents[2]


def _ocr_root() -> Path:
    return _ai_agents_root() / 'ocr'


def _legacy_site_packages() -> Path:
    return _ocr_root() / 'venv' / 'Lib' / 'site-packages'


def _extend_legacy_paths() -> None:
    for candidate in (_ocr_root(), _legacy_site_packages()):
        if candidate.exists():
            candidate_str = str(candidate)
            if candidate_str not in sys.path:
                sys.path.insert(0, candidate_str)


_extend_legacy_paths()

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PdfReader = None
    PYPDF2_AVAILABLE = False

try:
    import fitz  # type: ignore
    FITZ_AVAILABLE = True
except ImportError:
    fitz = None
    FITZ_AVAILABLE = False


class OCREngine:
    TEXT_EXTENSIONS = {'.txt', '.md', '.yaml', '.yml', '.ini', '.cfg', '.log', '.py', '.js', '.ts'}
    STRUCTURED_TEXT_EXTENSIONS = {'.json', '.csv', '.tsv', '.xml'}
    HTML_EXTENSIONS = {'.html', '.htm'}
    WORD_EXTENSIONS = {'.doc', '.docx'}
    SPREADSHEET_EXTENSIONS = {'.xls', '.xlsx'}
    RTF_EXTENSIONS = {'.rtf'}
    IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.webp', '.bmp', '.tiff'}
    MIN_TEXT_CHARS = 80

    def __init__(self, language: str = 'heb+eng'):
        self.language = language
        self.tessdata_dir = _ocr_root() / 'tessdata'
        self.tesseract_cmd = self._discover_tesseract_cmd()
        if self.tessdata_dir.exists():
            os.environ['TESSDATA_PREFIX'] = str(self.tessdata_dir)
        if TESSERACT_PYTHON_AVAILABLE and self.tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = self.tesseract_cmd

    def inspect(self, source_path: str) -> Dict[str, object]:
        path = Path(source_path)
        extension = path.suffix.lower()
        result = {
            'status': 'metadata_only',
            'text': '',
            'metadata': {
                'extension': extension,
                'size_bytes': path.stat().st_size if path.exists() else 0,
                'ocr_capabilities': self.capabilities(),
            },
        }
        if not path.exists():
            result['status'] = 'missing_file'
            return result
        if extension in self.TEXT_EXTENSIONS:
            text = path.read_text(encoding='utf-8', errors='ignore')
            result['status'] = 'text_extracted'
            result['text'] = text
            return result
        if extension in self.STRUCTURED_TEXT_EXTENSIONS:
            result.update(self._extract_structured_text(path))
            return result
        if extension in self.HTML_EXTENSIONS:
            raw = path.read_text(encoding='utf-8', errors='ignore')
            text = re.sub(r'<[^>]+>', ' ', raw)
            result['status'] = 'text_extracted'
            result['text'] = html.unescape(re.sub(r'\s+', ' ', text))
            return result
        if extension in self.RTF_EXTENSIONS:
            raw = path.read_text(encoding='utf-8', errors='ignore')
            result['status'] = 'text_extracted'
            result['text'] = self._extract_rtf_text(raw)
            return result
        if extension in self.WORD_EXTENSIONS:
            result.update(self._extract_word(path))
            return result
        if extension in self.SPREADSHEET_EXTENSIONS:
            result.update(self._extract_spreadsheet(path))
            return result
        if extension == '.epub':
            result.update(self._extract_epub(path))
            return result
        if extension == '.pdf':
            result.update(self._inspect_pdf(path))
            return result
        if extension in self.IMAGE_EXTENSIONS:
            result.update(self._inspect_image(path))
            return result
        fallback = self._inspect_generic_text(path)
        if fallback:
            result.update(fallback)
            return result
        return result

    def capabilities(self) -> Dict[str, object]:
        legacy_root = _ocr_root()
        legacy_packages = _legacy_site_packages()
        has_tessdata = self.tessdata_dir.exists()
        has_tesseract = bool(self.tesseract_cmd)
        return {
            'language': self.language,
            'tesseract_cmd': self.tesseract_cmd or '',
            'tesseract_available': has_tesseract,
            'tessdata_dir': str(self.tessdata_dir),
            'tessdata_available': has_tessdata,
            'pytesseract_available': TESSERACT_PYTHON_AVAILABLE,
            'pypdf2_available': PYPDF2_AVAILABLE,
            'fitz_available': FITZ_AVAILABLE,
            'pdf2image_available': PDF2IMAGE_AVAILABLE,
            'pil_available': Image is not None,
            'legacy_ocr_root': str(legacy_root),
            'legacy_ocr_root_available': legacy_root.exists(),
            'legacy_site_packages': str(legacy_packages),
            'legacy_site_packages_available': legacy_packages.exists(),
            'full_ocr_available': bool(has_tesseract and TESSERACT_PYTHON_AVAILABLE and FITZ_AVAILABLE and Image is not None),
            'text_extraction_available': bool(PYPDF2_AVAILABLE or FITZ_AVAILABLE),
        }

    def runtime_summary(self) -> Dict[str, object]:
        capabilities = self.capabilities()
        blockers: List[str] = []
        if not capabilities['tesseract_available']:
            blockers.append('tesseract_missing')
        if not capabilities['pytesseract_available']:
            blockers.append('pytesseract_missing')
        if not capabilities['fitz_available']:
            blockers.append('fitz_missing')
        return {
            'capabilities': capabilities,
            'ready_for_full_ocr': bool(capabilities['full_ocr_available']),
            'ready_for_text_extraction': bool(capabilities['text_extraction_available']),
            'blockers': blockers,
            'recommended_action': self.recommended_action(),
        }

    def recommended_action(self) -> str:
        capabilities = self.capabilities()
        if capabilities['full_ocr_available']:
            return 'להריץ OCR על הקבצים הממתינים.'
        if capabilities['text_extraction_available']:
            return 'להמשיך בחילוץ טקסט מובנה; עבור OCR מלא חסר כרגע Tesseract.'
        return 'להשלים סביבת OCR בסיסית לפני המשך עיבוד PDF.'

    def _discover_tesseract_cmd(self) -> str | None:
        env_value = os.environ.get('TESSERACT_CMD')
        candidates = [
            Path(env_value) if env_value else None,
            Path(r'C:\Program Files\Tesseract-OCR\tesseract.exe'),
            Path(r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe'),
            _ocr_root() / 'tesseract.exe',
        ]
        for candidate in candidates:
            if candidate and candidate.exists():
                return str(candidate)
        which_value = shutil.which('tesseract')
        if which_value:
            return which_value
        return None

    def _extract_structured_text(self, path: Path) -> Dict[str, object]:
        extension = path.suffix.lower()
        if extension == '.json':
            return self._extract_json_text(path)
        if extension in {'.csv', '.tsv'}:
            return self._extract_delimited_text(path)
        if extension == '.xml':
            raw = path.read_text(encoding='utf-8', errors='ignore')
            text = re.sub(r'<[^>]+>', ' ', raw)
            return {
                'status': 'text_extracted',
                'text': html.unescape(re.sub(r'\s+', ' ', text)).strip(),
                'metadata': {'parser': 'xml-tag-strip'},
            }
        return {'status': 'metadata_only', 'text': '', 'metadata': {'parser': 'structured-fallback'}}

    def _extract_json_text(self, path: Path) -> Dict[str, object]:
        try:
            payload = json.loads(path.read_text(encoding='utf-8-sig', errors='ignore'))
        except Exception as exc:
            return {'status': 'json_error', 'text': '', 'metadata': {'error': str(exc)}}

        lines: List[str] = []

        def _walk(value, prefix: str = '') -> None:
            if isinstance(value, dict):
                for key, child in value.items():
                    child_prefix = f'{prefix}.{key}' if prefix else str(key)
                    _walk(child, child_prefix)
                return
            if isinstance(value, list):
                for index, child in enumerate(value):
                    child_prefix = f'{prefix}[{index}]' if prefix else f'[{index}]'
                    _walk(child, child_prefix)
                return
            text = str(value).strip()
            if text:
                lines.append(f'{prefix}: {text}' if prefix else text)

        _walk(payload)
        rendered = '\n'.join(lines).strip()
        if not rendered:
            rendered = json.dumps(payload, ensure_ascii=False, indent=2)
        return {
            'status': 'text_extracted',
            'text': rendered,
            'metadata': {'parser': 'json-flattened', 'top_level_type': type(payload).__name__},
        }

    def _extract_delimited_text(self, path: Path) -> Dict[str, object]:
        delimiter = '\t' if path.suffix.lower() == '.tsv' else ','
        try:
            rows: List[str] = []
            with path.open('r', encoding='utf-8-sig', errors='ignore', newline='') as handle:
                reader = csv.reader(handle, delimiter=delimiter)
                for index, row in enumerate(reader):
                    if index > 250:
                        break
                    if not row or not any(str(cell).strip() for cell in row):
                        continue
                    rows.append(f'row {index + 1}: ' + ' | '.join(str(cell).strip() for cell in row))
            text = '\n'.join(rows).strip()
            return {
                'status': 'text_extracted' if text else 'metadata_only',
                'text': text,
                'metadata': {'parser': 'csv-reader', 'delimiter': delimiter},
            }
        except Exception as exc:
            return {'status': 'csv_error', 'text': '', 'metadata': {'error': str(exc)}}

    def _extract_word(self, path: Path) -> Dict[str, object]:
        extension = path.suffix.lower()
        if extension == '.docx':
            if PYTHON_DOCX_AVAILABLE:
                try:
                    document = python_docx.Document(str(path))
                    parts: List[str] = []
                    for paragraph in document.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            parts.append(text)
                    for table_index, table in enumerate(document.tables, start=1):
                        table_lines: List[str] = []
                        for row in table.rows:
                            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if cells:
                                table_lines.append(' | '.join(cells))
                        if table_lines:
                            parts.append(f'--- Table {table_index} ---')
                            parts.extend(table_lines)
                    text = '\n'.join(parts).strip()
                    return {
                        'status': 'text_extracted' if text else 'metadata_only',
                        'text': text,
                        'metadata': {'parser': 'python-docx'},
                    }
                except Exception as exc:
                    return {'status': 'docx_error', 'text': '', 'metadata': {'error': str(exc)}}
            return self._extract_docx_zip(path)
        if extension == '.doc':
            return self._extract_word_via_com(path)
        return {'status': 'metadata_only', 'text': '', 'metadata': {'error': f'unsupported word extension: {extension}'}}

    def _extract_docx_zip(self, path: Path) -> Dict[str, object]:
        try:
            with zipfile.ZipFile(path) as archive:
                xml = archive.read('word/document.xml').decode('utf-8', errors='ignore')
            text = re.sub(r'<[^>]+>', ' ', xml)
            return {
                'status': 'text_extracted',
                'text': re.sub(r'\s+', ' ', html.unescape(text)),
                'metadata': {'parser': 'zipfile-docx'},
            }
        except Exception as exc:
            return {'status': 'docx_error', 'text': '', 'metadata': {'error': str(exc)}}

    def _extract_word_via_com(self, path: Path) -> Dict[str, object]:
        if not WIN32COM_AVAILABLE or pythoncom is None or win32com is None:
            return {'status': 'doc_error', 'text': '', 'metadata': {'error': 'Microsoft Word COM unavailable'}}

        word_app = None
        document = None
        try:
            pythoncom.CoInitialize()
            word_app = win32com.client.DispatchEx('Word.Application')
            word_app.Visible = False
            word_app.DisplayAlerts = 0
            document = word_app.Documents.Open(str(path), ReadOnly=True, AddToRecentFiles=False, Visible=False)
            text = str(document.Range().Text or '').strip()
            text = re.sub(r'\n{3,}', '\n\n', text).strip()
            return {
                'status': 'text_extracted' if text else 'metadata_only',
                'text': text,
                'metadata': {'parser': 'word-com'},
            }
        except Exception as exc:
            return {'status': 'doc_error', 'text': '', 'metadata': {'error': str(exc), 'parser': 'word-com'}}
        finally:
            try:
                if document is not None:
                    document.Close(False)
            except Exception:
                pass
            try:
                if word_app is not None:
                    word_app.Quit()
            except Exception:
                pass
            try:
                if pythoncom is not None:
                    pythoncom.CoUninitialize()
            except Exception:
                pass

    def _extract_spreadsheet(self, path: Path) -> Dict[str, object]:
        extension = path.suffix.lower()
        if extension in {'.csv', '.tsv'}:
            return self._extract_delimited_text(path)

        com_result = self._extract_excel_via_com(path)
        if com_result.get('text'):
            return com_result

        if extension == '.xlsx':
            zip_result = self._extract_xlsx_zip(path)
            if zip_result.get('text'):
                return zip_result

        pandas_result = self._extract_excel_with_pandas(path)
        if pandas_result.get('text'):
            return pandas_result

        if com_result.get('metadata'):
            return com_result
        return pandas_result

    def _extract_excel_via_com(self, path: Path) -> Dict[str, object]:
        if not WIN32COM_AVAILABLE or pythoncom is None or win32com is None:
            return {'status': 'excel_error', 'text': '', 'metadata': {'error': 'Microsoft Excel COM unavailable'}}

        excel_app = None
        workbook = None
        try:
            pythoncom.CoInitialize()
            excel_app = win32com.client.DispatchEx('Excel.Application')
            excel_app.Visible = False
            excel_app.DisplayAlerts = False
            workbook = excel_app.Workbooks.Open(str(path), ReadOnly=True)
            snippets: List[str] = []
            for sheet in workbook.Worksheets:
                snippets.append(f'--- Sheet: {sheet.Name} ---')
                used_range = sheet.UsedRange
                rows = min(int(used_range.Rows.Count or 0), 250)
                cols = min(int(used_range.Columns.Count or 0), 40)
                for row_index in range(1, rows + 1):
                    values: List[str] = []
                    for col_index in range(1, cols + 1):
                        try:
                            cell_value = used_range.Cells(row_index, col_index).Value
                        except Exception:
                            cell_value = None
                        values.append('' if cell_value is None else str(cell_value).strip())
                    line = ' | '.join(values).strip()
                    if line:
                        snippets.append(line)
            text = '\n'.join(snippets).strip()
            return {
                'status': 'text_extracted' if text else 'metadata_only',
                'text': text,
                'metadata': {'parser': 'excel-com'},
            }
        except Exception as exc:
            return {'status': 'excel_error', 'text': '', 'metadata': {'error': str(exc), 'parser': 'excel-com'}}
        finally:
            try:
                if workbook is not None:
                    workbook.Close(False)
            except Exception:
                pass
            try:
                if excel_app is not None:
                    excel_app.Quit()
            except Exception:
                pass
            try:
                if pythoncom is not None:
                    pythoncom.CoUninitialize()
            except Exception:
                pass

    def _extract_excel_with_pandas(self, path: Path) -> Dict[str, object]:
        if not PANDAS_AVAILABLE:
            return {'status': 'excel_error', 'text': '', 'metadata': {'error': 'pandas unavailable'}}
        try:
            sheets = pd.read_excel(str(path), sheet_name=None)  # type: ignore[union-attr]
            snippets: List[str] = []
            for sheet_name, frame in (sheets or {}).items():
                snippets.append(f'--- Sheet: {sheet_name} ---')
                if frame is None or frame.empty:
                    continue
                header = [str(col) for col in frame.columns.tolist()]
                if header:
                    snippets.append(' | '.join(header))
                for _, row in frame.head(200).iterrows():
                    snippets.append(' | '.join('' if cell is None else str(cell).strip() for cell in row.tolist()))
            text = '\n'.join(snippets).strip()
            return {
                'status': 'text_extracted' if text else 'metadata_only',
                'text': text,
                'metadata': {'parser': 'pandas-excel'},
            }
        except Exception as exc:
            return {'status': 'excel_error', 'text': '', 'metadata': {'error': str(exc), 'parser': 'pandas-excel'}}

    def _extract_xlsx_zip(self, path: Path) -> Dict[str, object]:
        try:
            with zipfile.ZipFile(path) as archive:
                shared_strings = self._read_xlsx_shared_strings(archive)
                sheet_map = self._read_xlsx_sheet_map(archive)
                if not sheet_map:
                    return {'status': 'excel_error', 'text': '', 'metadata': {'error': 'workbook sheets missing'}}

                snippets: List[str] = []
                for sheet_name, sheet_path in sheet_map.items():
                    if not sheet_path:
                        continue
                    try:
                        raw = archive.read(sheet_path)
                    except Exception:
                        continue
                    sheet_text = self._parse_xlsx_sheet(raw, shared_strings)
                    if sheet_text.strip():
                        snippets.append(f'--- Sheet: {sheet_name} ---')
                        snippets.append(sheet_text)
                text = '\n'.join(snippets).strip()
                return {
                    'status': 'text_extracted' if text else 'metadata_only',
                    'text': text,
                    'metadata': {'parser': 'xlsx-zip'},
                }
        except Exception as exc:
            return {'status': 'excel_error', 'text': '', 'metadata': {'error': str(exc), 'parser': 'xlsx-zip'}}

    def _read_xlsx_shared_strings(self, archive: zipfile.ZipFile) -> List[str]:
        try:
            data = archive.read('xl/sharedStrings.xml')
        except Exception:
            return []
        try:
            root = ET.fromstring(data)
        except Exception:
            return []
        ns = {'a': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
        strings: List[str] = []
        for si in root.findall('a:si', ns):
            parts = [node.text or '' for node in si.findall('.//a:t', ns)]
            strings.append(''.join(parts))
        return strings

    def _read_xlsx_sheet_map(self, archive: zipfile.ZipFile) -> Dict[str, str]:
        try:
            workbook_xml = ET.fromstring(archive.read('xl/workbook.xml'))
        except Exception:
            return {}
        try:
            rels_xml = ET.fromstring(archive.read('xl/_rels/workbook.xml.rels'))
        except Exception:
            rels_xml = None
        workbook_ns = {
            'a': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        }
        rels_map: Dict[str, str] = {}
        if rels_xml is not None:
            for rel in rels_xml:
                rel_id = rel.attrib.get('Id')
                rel_target = rel.attrib.get('Target', '')
                if rel_id and rel_target:
                    rels_map[rel_id] = rel_target.lstrip('/')
        sheet_map: Dict[str, str] = {}
        for sheet in workbook_xml.findall('a:sheets/a:sheet', workbook_ns):
            name = sheet.attrib.get('name', 'Sheet')
            rel_id = sheet.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', '')
            target = rels_map.get(rel_id, '')
            if target and not target.startswith('xl/'):
                target = f'xl/{target}'
            sheet_map[name] = target
        return sheet_map

    def _parse_xlsx_sheet(self, raw_xml: bytes, shared_strings: List[str]) -> str:
        try:
            root = ET.fromstring(raw_xml)
        except Exception:
            return ''
        ns = {'a': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
        rows: List[str] = []
        for row in root.findall('.//a:sheetData/a:row', ns):
            cells: Dict[int, str] = {}
            max_index = 0
            for cell in row.findall('a:c', ns):
                ref = cell.attrib.get('r', '')
                col_index = self._xlsx_column_index(ref)
                if col_index <= 0:
                    continue
                max_index = max(max_index, col_index)
                cell_type = cell.attrib.get('t', '')
                value = ''
                if cell_type == 's':
                    shared_index = cell.findtext('a:v', default='', namespaces=ns)
                    if shared_index.isdigit():
                        idx = int(shared_index)
                        if 0 <= idx < len(shared_strings):
                            value = shared_strings[idx]
                elif cell_type == 'inlineStr':
                    value = ''.join(node.text or '' for node in cell.findall('.//a:t', ns))
                else:
                    value = cell.findtext('a:v', default='', namespaces=ns)
                cells[col_index] = re.sub(r'\s+', ' ', str(value or '').strip())
            if max_index == 0:
                continue
            ordered = [cells.get(index, '') for index in range(1, max_index + 1)]
            line = ' | '.join(ordered).strip()
            if line:
                rows.append(line)
        return '\n'.join(rows)

    def _xlsx_column_index(self, cell_ref: str) -> int:
        letters = ''.join(ch for ch in str(cell_ref or '') if ch.isalpha()).upper()
        if not letters:
            return 0
        value = 0
        for char in letters:
            value = value * 26 + (ord(char) - 64)
        return value

    def _extract_epub(self, path: Path) -> Dict[str, object]:
        try:
            snippets: List[str] = []
            with zipfile.ZipFile(path) as archive:
                members = [name for name in archive.namelist() if name.lower().endswith(('.xhtml', '.html', '.htm', '.xml'))]
                for name in members[:25]:
                    raw = archive.read(name).decode('utf-8', errors='ignore')
                    text = re.sub(r'<[^>]+>', ' ', raw)
                    clean = html.unescape(re.sub(r'\s+', ' ', text)).strip()
                    if clean:
                        snippets.append(clean)
            joined = '\n'.join(snippets)
            return {
                'status': 'text_extracted' if joined else 'metadata_only',
                'text': joined,
                'metadata': {'parser': 'zipfile-epub', 'documents_sampled': min(len(snippets), 25)},
            }
        except Exception as exc:
            return {'status': 'epub_error', 'text': '', 'metadata': {'error': str(exc)}}

    def _extract_rtf_text(self, raw: str) -> str:
        text = re.sub(r'\\par[d]?|\\line', '\n', raw)
        text = re.sub(r'\\[a-zA-Z]+-?\d* ?', ' ', text)
        text = re.sub(r'\\[{}\\]', ' ', text)
        text = re.sub(r'{\\[^{}]*}|[{}]', ' ', text)
        text = re.sub(r'\\u-?\d+\??', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        return html.unescape(text).strip()

    def _inspect_generic_text(self, path: Path) -> Dict[str, object] | None:
        try:
            data = path.read_bytes()
        except Exception:
            return None
        if not data:
            return None
        sample = data[:8192]
        if self._looks_binary(sample):
            return None
        for encoding in ('utf-8', 'utf-8-sig', 'cp1255', 'cp1252', 'latin-1'):
            try:
                text = data.decode(encoding, errors='strict')
                break
            except Exception:
                text = ''
        if not text.strip():
            return None
        return {
            'status': 'text_extracted',
            'text': re.sub(r'\s+', ' ', text).strip(),
            'metadata': {'parser': f'generic-text:{path.suffix.lower() or "unknown"}'},
        }

    def _looks_binary(self, sample: bytes) -> bool:
        if not sample:
            return False
        if b'\x00' in sample:
            return True
        printable = sum(1 for byte in sample if byte in b'\n\r\t\f\b' or 32 <= byte <= 126)
        return (printable / len(sample)) < 0.7

    def _inspect_pdf(self, path: Path) -> Dict[str, object]:
        py_pdf_result = self._extract_pdf_text_with_pypdf2(path)
        py_pdf_text = str(py_pdf_result.get('text') or '')
        if len(py_pdf_text.strip()) >= self.MIN_TEXT_CHARS:
            return py_pdf_result

        fitz_result = self._extract_pdf_text_with_fitz(path)
        fitz_text = str(fitz_result.get('text') or '')
        if len(fitz_text.strip()) > len(py_pdf_text.strip()):
            candidate = fitz_result
        else:
            candidate = py_pdf_result
        if len(str(candidate.get('text') or '').strip()) >= self.MIN_TEXT_CHARS:
            return candidate

        if self._can_run_ocr() and FITZ_AVAILABLE:
            ocr_result = self._ocr_pdf_with_fitz(path)
            if len((ocr_result.get('text') or '').strip()) >= self.MIN_TEXT_CHARS:
                return ocr_result

        text = str(candidate.get('text') or '').strip()
        candidate['status'] = 'ocr_pending' if not text else 'text_extracted'
        return candidate

    def _extract_pdf_text_with_pypdf2(self, path: Path) -> Dict[str, object]:
        if not PYPDF2_AVAILABLE:
            return {'status': 'ocr_pending', 'text': '', 'metadata': {'parser': 'pdf-no-parser'}}
        try:
            reader = PdfReader(str(path))
            parts: List[str] = []
            total_pages = len(reader.pages)
            for index, page in enumerate(reader.pages[:40], start=1):
                try:
                    text = page.extract_text() or ''
                except Exception:
                    text = ''
                if text.strip():
                    parts.append(f'--- Page {index} ---\n{text.strip()}')
            joined = '\n'.join(parts)
            return {
                'status': 'text_extracted' if joined.strip() else 'ocr_pending',
                'text': joined,
                'metadata': {'parser': 'PyPDF2', 'pages_sampled': min(total_pages, 40)},
            }
        except Exception as exc:
            return {'status': 'pdf_error', 'text': '', 'metadata': {'error': str(exc), 'parser': 'PyPDF2'}}

    def _extract_pdf_text_with_fitz(self, path: Path) -> Dict[str, object]:
        if not FITZ_AVAILABLE:
            return {'status': 'ocr_pending', 'text': '', 'metadata': {'parser': 'fitz-unavailable'}}
        try:
            document = fitz.open(str(path))
            parts: List[str] = []
            total_pages = len(document)
            for index, page in enumerate(document[:40], start=1):
                try:
                    text = page.get_text('text') or ''
                except Exception:
                    text = ''
                if text.strip():
                    parts.append(f'--- Page {index} ---\n{text.strip()}')
            joined = '\n'.join(parts)
            return {
                'status': 'text_extracted' if joined.strip() else 'ocr_pending',
                'text': joined,
                'metadata': {'parser': 'fitz-text', 'pages_sampled': min(total_pages, 40)},
            }
        except Exception as exc:
            return {'status': 'pdf_error', 'text': '', 'metadata': {'error': str(exc), 'parser': 'fitz-text'}}

    def _ocr_pdf_with_fitz(self, path: Path) -> Dict[str, object]:
        try:
            document = fitz.open(str(path))
            snippets: List[str] = []
            for index, page in enumerate(document[:3], start=1):
                pix = page.get_pixmap(dpi=180)
                image = Image.frombytes('RGB', (pix.width, pix.height), pix.samples)
                snippets.append(pytesseract.image_to_string(image, lang=self.language))
            text = '\n'.join(part.strip() for part in snippets if part.strip())
            return {
                'status': 'ocr_extracted' if text else 'ocr_pending',
                'text': text,
                'metadata': {'parser': 'fitz+tesseract', 'pages_sampled': min(len(document), 3)},
            }
        except Exception as exc:
            return {'status': 'ocr_error', 'text': '', 'metadata': {'error': str(exc)}}

    def _inspect_image(self, path: Path) -> Dict[str, object]:
        if self._can_run_ocr() and Image is not None:
            try:
                text = pytesseract.image_to_string(Image.open(path), lang=self.language).strip()
                return {
                    'status': 'ocr_extracted' if text else 'ocr_empty',
                    'text': text,
                    'metadata': {'parser': 'tesseract-image'},
                }
            except Exception as exc:
                return {'status': 'ocr_error', 'text': '', 'metadata': {'error': str(exc)}}
        return {'status': 'ocr_pending', 'text': '', 'metadata': {'parser': 'image-metadata-only'}}

    def _can_run_ocr(self) -> bool:
        return bool(TESSERACT_PYTHON_AVAILABLE and self.tesseract_cmd and FITZ_AVAILABLE and Image is not None)
