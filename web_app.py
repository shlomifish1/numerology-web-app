
import streamlit as st
import datetime
import os
import sys
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add current directory to path so imports work correctly if needed
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from numerology_calculator import NumerologyCalculator
from config_manager import ConfigManager
from gpt_report import generate_person_report
from research.active_book_lab import render_lab_panel, render_live_books_panel
import personal_y

# --- Page Config ---
st.set_page_config(
    page_title="מפה נומרולוגית מקצועית 🔮 | ניתוח אישי מעמיק",
    page_icon="ðŸ”®",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'About': "מערכת מקצועית לניתוח נומרולוגי מעמיק. מבוסס על חכמת המספרים העתיקה."
    }
)

# --- Load Custom CSS ---
def load_css():
    css_file = os.path.join(os.path.dirname(__file__), "static", "custom.css")
    if os.path.exists(css_file):
        with open(css_file, encoding='utf-8') as f:
            st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)
    else:
        # Fallback basic RTL if custom CSS not found
        st.markdown("""
        <style>
            .stApp {
                direction: rtl;
                text-align: right;
            }
            h1, h2, h3, h4, h5, h6, p, div {
                text-align: right;
            }
        </style>
        """, unsafe_allow_html=True)

load_css()

def _set_rtl_paragraph(paragraph):
    """ Helper function to set paragraph to RTL and right-aligned. """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.rtl = True


def _export_interpretation(calc, category, number, gender_key, **kwargs):
    return calc.render_interpretation(category, number, gender_key, **kwargs)

def generate_docx_bytes(calc, gender_key, ai_report_text=None):
    """Generates the Word document in memory and returns a BytesIO object."""
    document = Document()
    
    # Title
    heading = document.add_heading(f'דוח נומרולוגי עבור: {calc.full_name}', level=0)
    _set_rtl_paragraph(heading)
    
    # Info
    p = document.add_paragraph(f"תאריך לידה: {calc.p_day}/{calc.p_month}/{calc.p_year} (מקורי: {calc.full_date_short})") 
    _set_rtl_paragraph(p)
    p = document.add_paragraph(f"מין: {'זכר' if gender_key == 'male' else 'נקבה'}")
    _set_rtl_paragraph(p)
    p = document.add_paragraph(f"תאריך הפקת הדוח: {datetime.date.today().strftime('%d/%m/%Y')}")
    _set_rtl_paragraph(p)
    document.add_paragraph()

    # AI Report Section (if available)
    if ai_report_text:
        h = document.add_heading('המלצת עומק (AI)', level=1)
        _set_rtl_paragraph(h)
        p = document.add_paragraph(ai_report_text)
        _set_rtl_paragraph(p)
        document.add_paragraph()

    # Destiny
    if calc.final_number_destiny is not None:
        h = document.add_heading('מספר ייעוד (דרך חיים)', level=1)
        _set_rtl_paragraph(h)
        p = document.add_paragraph(f"המספר המחושב: {calc.final_number_destiny}")
        _set_rtl_paragraph(p)
        interpretation = _export_interpretation(calc, "destiny", calc.final_number_destiny, gender_key)
        p = document.add_paragraph(interpretation)
        _set_rtl_paragraph(p)
        document.add_paragraph()

    # Personal Year
    if calc.shana_ishit is not None:
        h = document.add_heading('שנה אישית', level=1)
        _set_rtl_paragraph(h)
        p = document.add_paragraph(f"השנה האישית שלך: {calc.shana_ishit}")
        _set_rtl_paragraph(p)
        interpretation = _export_interpretation(calc, "personal_year", calc.shana_ishit, gender_key)
        p = document.add_paragraph(interpretation)
        _set_rtl_paragraph(p)
        document.add_paragraph()
    
    # Birth Day
    if calc.p_day is not None:
        h = document.add_heading('יום לידה', level=1)
        _set_rtl_paragraph(h)
        p = document.add_paragraph(f"יום הלידה המצומצם: {calc.p_day}")
        _set_rtl_paragraph(p)
        interpretation = _export_interpretation(calc, "birth_day", calc.p_day, gender_key)
        p = document.add_paragraph(interpretation)
        _set_rtl_paragraph(p)
        document.add_paragraph()

    # Peaks and Challenges
    h_main_peaks = document.add_heading('פסגות ואתגרים', level=1)
    _set_rtl_paragraph(h_main_peaks)
    
    peaks_data = [
        ("תקופה ראשונה", calc.peak1_reduced, calc.challenge1_reduced, f"עד גיל {calc.first_pick_start}"),
        ("תקופה שנייה", calc.peak2_reduced, calc.challenge2_reduced, f"מגיל {calc.first_pick_start + 1} עד {calc.second_pick_start - 1}"),
        ("תקופה שלישית", calc.peak3_reduced, calc.challenge3_reduced, f"מגיל {calc.second_pick_start} עד {calc.third_pick_start - 1}"),
        ("תקופה רביעית", calc.peak4_reduced, calc.challenge4_reduced, f"מגיל {calc.third_pick_start} ואילך")
    ]
    
    for period_name, peak, challenge, age_range in peaks_data:
        h_period = document.add_heading(f"{period_name} ({age_range})", level=2)
        _set_rtl_paragraph(h_period)
        
        p = document.add_paragraph(f"פסגה: {peak}")
        _set_rtl_paragraph(p)
        interp_peak = _export_interpretation(calc, "peaks_interpretation", peak, gender_key)
        p = document.add_paragraph(interp_peak)
        _set_rtl_paragraph(p)
        
        p = document.add_paragraph(f"אתגר: {challenge}")
        _set_rtl_paragraph(p)
        interp_challenge = _export_interpretation(calc, "challenges_interpretation", challenge, gender_key)
        p = document.add_paragraph(interp_challenge)
        _set_rtl_paragraph(p)
        
        document.add_paragraph()

    # Save to IO
    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio


def _render_runtime_interpretation(calc, category, value, gender_key, **kwargs):
    if value in (None, ""):
        st.info("אין ערך מחושב להצגה.")
        return
    text = calc.get_interpretation(category, value, gender_key, **kwargs)
    if isinstance(text, str) and "קובץ פרשנות לא נמצא" in text:
        st.warning(f"חסר קובץ פרשנות עבור `{category}` בערך `{value}`.")
        st.caption(text)
    else:
        st.write(text)


def _render_peak_challenge_grid(calc, gender_key):
    periods = [
        ("פסגה ראשונה", calc.peak1_reduced, "אתגר ראשון", calc.challenge1_reduced, f"עד גיל {calc.first_pick_start}"),
        ("פסגה שניה", calc.peak2_reduced, "אתגר שני", calc.challenge2_reduced, f"מגיל {calc.first_pick_start + 1} עד {calc.second_pick_start - 1}"),
        ("פסגה שלישית", calc.peak3_reduced, "אתגר שלישי", calc.challenge3_reduced, f"מגיל {calc.second_pick_start} עד {calc.third_pick_start - 1}"),
        ("פסגה רביעית", calc.peak4_reduced, "אתגר רביעי", calc.challenge4_reduced, f"מגיל {calc.third_pick_start} ואילך"),
    ]

    for index, (peak_label, peak_value, challenge_label, challenge_value, age_label) in enumerate(periods, start=1):
        combo_key = f"{peak_value}{challenge_value}" if peak_value is not None and challenge_value is not None else ""
        st.markdown(f"#### תקופה {index} · {age_label}")
        col_peak, col_challenge, col_combo = st.columns(3)
        with col_peak:
            st.metric(peak_label, peak_value)
            with st.expander(f"פירוש {peak_label}", expanded=False):
                _render_runtime_interpretation(calc, "peaks_interpretation", peak_value, gender_key)
        with col_challenge:
            st.metric(challenge_label, challenge_value)
            with st.expander(f"פירוש {challenge_label}", expanded=False):
                _render_runtime_interpretation(calc, "challenges_interpretation", challenge_value, gender_key)
        with col_combo:
            st.metric("שילוב פסגה ואתגר", combo_key or "-")
            with st.expander("פירוש שילוב", expanded=False):
                if not combo_key:
                    st.info("אין שילוב מחושב לתקופה זו.")
                else:
                    combo_text = calc.get_interpretation(
                        "peak_challenge_comb",
                        combo_key,
                        gender_key,
                        is_peak_challenge_comb=True,
                    )
                    if isinstance(combo_text, str) and "קובץ פרשנות לא נמצא" in combo_text:
                        st.warning(f"חסר שילוב מדויק עבור פסגה {peak_value} ואתגר {challenge_value}.")
                        st.caption(combo_text)
                    else:
                        st.write(combo_text)
        st.divider()

def main():
    # Initialize Config
    config_manager = ConfigManager()
    
    # Initialize Session State
    if "calc_done" not in st.session_state:
        st.session_state["calc_done"] = False
    if "calc_obj" not in st.session_state:
        st.session_state["calc_obj"] = None
    if "gender_key" not in st.session_state:
        st.session_state["gender_key"] = None
    if "ai_report_text" not in st.session_state:
        st.session_state["ai_report_text"] = None
    if "user_email" not in st.session_state:
        st.session_state["user_email"] = ""

    # --- Hero Section (shown when no calculation yet) ---
    if not st.session_state.get("calc_done"):
        st.markdown("""
        <div style='text-align: center; padding: 2rem 1rem;'>
            <h1 style='font-size: 3rem; margin-bottom: 1rem;'>🔮 המפה הנומרולוגית שלך מחכה</h1>
            <p style='font-size: 1.3rem; color: #CBD5E1; max-width: 700px; margin: 0 auto 2rem auto;'>
                גלה את ייעודך האמיתי, החוזקות הנסתרות שלך, והתובנות המנחות לעת ידך הנוכחי.
                מבוסס על חכמת המספרים העתיקה המשולבת עם AI מתקדם.
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # Value Proposition
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown("""
            ### 🎯 ניתוח מדויק
            חישובים נומרולוגיים מקצועיים המבוססים על שמך המלא ותאריך הלידה
            """)
        with col2:
            st.markdown("""
            ### 🤖 תובנות AI
            המלצות מותאמות אישית על ידי בינה מלאכותית מתקדמת
            """)
        with col3:
            st.markdown("""
            ### 📊 דוח מקצועי
            מסמך Word מעוצב ומפורט עם כל הפרשנויות
            """)

    # --- Sidebar Inputs ---
    with st.sidebar:
        st.header("📝 הזן את הפרטים שלך")
        
        first_name = st.text_input("שם פרטי", placeholder="לדוגמה: דוד")
        last_name = st.text_input("שם משפחה", placeholder="לדוגמה: כהן")
        
        # Date Input: Starts empty (None) for better UX
        dob = st.date_input(
            "תאריך לידה",
            min_value=datetime.date(1920, 1, 1),
            max_value=datetime.date.today(),
            value=None,  # Empty initial value
            format="DD/MM/YYYY" 
        )
        
        gender = st.selectbox("מין", ["זכר", "נקבה"])
        
        # Email input for order processing
        user_email = st.text_input("כתובת Email (לקבלת הדוח)", 
                                    value=st.session_state.get("user_email", ""),
                                    placeholder="example@email.com")
        if user_email:
            st.session_state["user_email"] = user_email
        
        st.markdown("---")
        calculate_btn = st.button("🔍 הצג תצוגה מקדימה", type="primary", use_container_width=True)

    # --- Calculations ---
    if calculate_btn:
        if not first_name or not last_name or dob is None:
            st.error("⚠️ אנא מלא את כל השדות: שם פרטי, שם משפחה ותאריך לידה")
        else:
            # Prepare data for calculator
            day_str = str(dob.day).zfill(2)
            month_str = str(dob.month).zfill(2)
            year_str = str(dob.year)
            gender_key = "male" if gender == "זכר" else "female"

            # Initialize Calculator
            calc = NumerologyCalculator()
            
            try:
                with st.spinner("🔮 מחשב את המפה הנומרולוגית שלך..."):
                    calc.calculate(day_str, month_str, year_str, first_name, last_name, gender_key)
                
                # Save to Session State
                st.session_state["calc_done"] = True
                st.session_state["calc_obj"] = calc
                st.session_state["gender_key"] = gender_key
                st.session_state["ai_report_text"] = None # Reset AI report on new calculation
                st.rerun()
                
            except Exception as e:
                st.error(f"❌ אירעה שגיאה בעת חישוב הנתונים: {e}")
                st.exception(e)
                st.session_state["calc_done"] = False

    # --- Display Results ---
    if st.session_state.get("calc_done"):
        calc = st.session_state["calc_obj"]
        gender_key = st.session_state["gender_key"]
        
        # Success message with name
        st.markdown(f"""
        <div style='background: linear-gradient(135deg, rgba(16, 185, 129, 0.15), rgba(20, 184, 166, 0.15)); 
                    padding: 1.5rem; border-radius: 12px; border-right: 4px solid #10B981; margin-bottom: 2rem;'>
            <h2 style='color: #D1FAE5; margin: 0;'>✨ המפה הנומרולוגית של {calc.full_name}</h2>
            <p style='color: #CCFBF1; margin: 0.5rem 0 0 0;'>להלן תצוגה מקדימה של החישובים הבסיסיים שלך</p>
        </div>
        """, unsafe_allow_html=True)
        
        # --- Pricing / Order Section (at the top) ---
        st.markdown("### 💎 הזמן את המפה המלאה שלך")
        st.markdown("תצוגה זו כוללת רק חלק מהניתוח. קבל את הדוח המלא במייל תוך 24 שעות:")
        
        col_short, col_long = st.columns(2)
        
        with col_short:
            st.markdown("""
            <div style='background: rgba(107, 70, 193, 0.15); border: 2px solid #6B46C1; border-radius: 16px; padding: 2rem; text-align: center;'>
                <h3 style='color: #E9D5FF; margin-top: 0;'>🌟 מפה קצרה</h3>
                <p style='font-size: 2.5rem; font-weight: bold; color: #14B8A6; margin: 1rem 0;'>₪99</p>
                <ul style='text-align: right; color: #CBD5E1; list-style: none; padding: 0;'>
                    <li>✅ מספר ייעוד ופסגות</li>
                    <li>✅ שנה אישית ויום לידה</li>
                    <li>✅ דוח Word מעוצב</li>
                    <li>✅ משלוח במייל תוך 24 שעות</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("🛒 הזמן מפה קצרה", use_container_width=True, type="primary"):
                st.session_state["order_type"] = "short"
                st.markdown(f"""
                <meta http-equiv="refresh" content="0; url=https://www.meshulam.co.il/purchase?b=dd556e28193228493a6cf2a4568b7cbd">
                <p style='text-align: center; color: #14B8A6;'>מעביר אותך לעמוד התשלום... ⏳</p>
                """, unsafe_allow_html=True)
        
        with col_long:
            st.markdown("""
            <div style='background: rgba(20, 184, 166, 0.15); border: 2px solid #14B8A6; border-radius: 16px; padding: 2rem; text-align: center;'>
                <h3 style='color: #CCFBF1; margin-top: 0;'>💎 מפה ארוכה ומעמיקה</h3>
                <p style='font-size: 2.5rem; font-weight: bold; color: #F59E0B; margin: 1rem 0;'>₪299</p>
                <ul style='text-align: right; color: #CBD5E1; list-style: none; padding: 0;'>
                    <li>✅ כל מה שבמפה הקצרה +</li>
                    <li>✅ ריבוע פיתגורס מלא</li>
                    <li>✅ ניתוח רבעונים</li>
                    <li>✅ תובנות AI מתקדמות</li>
                    <li>✅ דוח מקצועי בן 10+ עמודים</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("🛒 הזמן מפה ארוכה", use_container_width=True, type="primary"):
                st.session_state["order_type"] = "long"
                st.markdown(f"""
                <meta http-equiv="refresh" content="0; url=https://meshulam.co.il/purchase?b=19e85446f85d006f13a4a0b963a8b79d">
                <p style='text-align: center; color: #14B8A6;'>מעביר אותך לעמוד התשלום... ⏳</p>
                """, unsafe_allow_html=True)
        
        st.markdown("---")

        
        # --- AI Section ---
        with st.expander("✨ שאל את ה-AI (המלצה אישית)", expanded=False):
            if st.button("קבל המלצת עומק מה-AI"):
                api_key = config_manager.get_api_key("google_ai")
                if not api_key:
                    st.error("לא נמצא מפתח API חוקי בהגדרות. אנא בדוק את הקובץ config.json.")
                else:
                    with st.spinner("ה-AI מנתח את כל המפה... זה עשוי לקחת רגע..."):
                        # Prepare Data Dict
                        ai_data = {
                            "full_name": calc.full_name,
                            "birth_date": f"{calc.original_day:02d}{calc.original_month:02d}{calc.original_year}",
                            "personal_day": calc.p_day,
                            "personal_month": calc.p_month,
                            "personal_year": calc.p_year,
                            "destiny_number": calc.final_number_destiny,
                            "personal_year_number": calc.shana_ishit,
                            "hidden_year": calc.shana_nisteret,
                            "age": calc.age,
                            "life_peaks": [calc.peak1_reduced, calc.peak2_reduced, calc.peak3_reduced, calc.peak4_reduced],
                            "challenges": [calc.challenge1_reduced, calc.challenge2_reduced, calc.challenge3_reduced, calc.challenge4_reduced],
                            "quarters": [calc.first_quarter_reduced, calc.second_quarter_reduced, calc.third_quarter_reduced, calc.forth_quarter_reduced],
                            "gender": "זכר" if gender_key == "male" else "נקבה",
                            "green_context": calc.calc_green_snapshot(),
                        }
                        # Call AI
                        report = generate_person_report(ai_data, model_name=config_manager.get("active_model", "gemini-flash-latest"), api_key=api_key)
                        st.session_state["ai_report_text"] = report
        
        # Display cached AI report if exists
        if st.session_state["ai_report_text"]:
            st.info(st.session_state["ai_report_text"])

        # --- Export Options (Top) ---
        col_dl, col_blank = st.columns([1, 4])
        with col_dl:
                # Pass AI text to DOCX generator
                docx_file = generate_docx_bytes(calc, gender_key, ai_report_text=st.session_state["ai_report_text"])
                st.download_button(
                    label="📄 הורד דוח Word",
                    data=docx_file,
                    file_name=f"Numerology_Report_{calc.full_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )

        map_tab, lab_tab, books_tab = st.tabs(["??? ?????", "?????", "?????"])

        with map_tab:
            st.caption("???? ?????? ????? ?? ???? ???? ???? men/women ?-runtime.")
            main_tab, peaks_tab, meanings_tab = st.tabs(["??????", "????? ????? ??????", "????? ??????? ????? ??????"])

            with main_tab:
                col1, col2, col3 = st.columns(3)

                with col1:
                    st.image("https://img.icons8.com/color/96/numerology.png", width=64)
                    st.metric("???? ?????", calc.final_number_destiny)
                    with st.expander("????? ???? ?????"):
                        _render_runtime_interpretation(calc, "destiny", calc.final_number_destiny, gender_key)

                with col2:
                    st.metric("??? ?????", calc.shana_ishit)
                    with st.expander("????? ??? ?????"):
                        _render_runtime_interpretation(calc, "personal_year", calc.shana_ishit, gender_key)

                with col3:
                    st.metric("??? ????", calc.p_day)
                    with st.expander("????? ??? ????"):
                        st.write(f"??? ????/?????: {calc.p_day}")
                        _render_runtime_interpretation(calc, "birth_day", calc.p_day, gender_key)

                st.markdown("### ??? ?????")
                c1, c2, c3, c4 = st.columns(4)
                c1.metric("?? ????", calc.first_name_val)
                c2.metric("?? ???", calc.full_name_val)
                c3.metric("????? ??????", calc.itzurim_val)
                c4.metric("??????", calc.aiv_val)

            with peaks_tab:
                st.subheader("?????, ?????? ????? ????? ??????")
                _render_peak_challenge_grid(calc, gender_key)

            with meanings_tab:
                st.info("??? ?????? ????? ????? ?????? ???? ???? ?-runtime ????.")
                if calc.shana_nisteret:
                    st.write(f"**??? ?????:** {calc.shana_nisteret}")
                    _render_runtime_interpretation(
                        calc,
                        "hidden_year",
                        str(calc.shana_nisteret),
                        gender_key,
                        is_hidden_year=True,
                    )

        with lab_tab:
            st.caption("?????? ????? ?? ???? ???? ???? ???? interpretations/research.")
            render_lab_panel(prefix="web_lab", calc=calc)

        with books_tab:
            st.caption("?????? ?????? ?????? ????? ????? ?????? ????? ??????.")
            render_live_books_panel(prefix="web_books", calc=calc)

    else:
        st.info("הזן פרטים בצד ימין ולחץ על 'בצע חישוב' כדי להתחיל.")

if __name__ == "__main__":
    main()


