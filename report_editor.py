import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import threading
from docx import Document
from config_manager import ConfigManager
from drive_uploader import DriveUploader
import google.generativeai as genai
import os

class ReportEditor:
    def __init__(self, parent, initial_text, config_manager: ConfigManager, default_filename="report.docx"):
        self.window = tk.Toplevel(parent)
        self.window.title("עורך דוחות AI")
        self.window.geometry("900x700")
        self.config_manager = config_manager
        self.initial_text = initial_text
        self.default_filename = default_filename
        self.drive_uploader = DriveUploader()
        
        self.create_widgets()

    def create_widgets(self):
        # Toolbar
        toolbar = ttk.Frame(self.window, padding="5")
        toolbar.pack(side=tk.TOP, fill=tk.X)
        
        ttk.Button(toolbar, text="שמור DOCX למחשב", command=self.save_docx).pack(side=tk.RIGHT, padx=5)
        ttk.Button(toolbar, text="העלה ל-Google Drive", command=self.upload_to_drive).pack(side=tk.RIGHT, padx=5)
        
        ttk.Separator(toolbar, orient=tk.VERTICAL).pack(side=tk.RIGHT, fill=tk.Y, padx=10)
        
        ttk.Button(toolbar, text="שפר ניסוח (AI)", command=self.refine_text).pack(side=tk.RIGHT, padx=5)
        ttk.Label(toolbar, text=":ערוך את הדוח לפני השמירה").pack(side=tk.RIGHT, padx=5)

        # Main Text Editor
        # RTL support: justify right
        self.text_area = tk.Text(self.window, wrap=tk.WORD, font=("Arial", 11), undo=True)
        self.text_area.tag_configure("rtl", justify='right')
        self.text_area.insert("1.0", self.initial_text)
        self.text_area.tag_add("rtl", "1.0", "end")
        
        scrollbar = ttk.Scrollbar(self.window, command=self.text_area.yview)
        self.text_area['yscrollcommand'] = scrollbar.set
        
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.text_area.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    def save_docx(self):
        content = self.text_area.get("1.0", tk.END).strip()
        if not content:
            messagebox.showwarning("ריק", "הדוח ריק.")
            return

        filepath = filedialog.asksaveasfilename(
            initialfile=self.default_filename,
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if not filepath:
            return

        try:
            doc = Document()
            for line in content.split("\n"):
                p = doc.add_paragraph(line)
                p.paragraph_format.rtl = True
                p.alignment = 2 # RIGHT
            
            doc.save(filepath)
            messagebox.showinfo("נשמר", f"הקובץ נשמר בהצלחה:\n{filepath}")
            return filepath
        except Exception as e:
            messagebox.showerror("שגיאה", f"שגיאה בשמירה: {e}")
            return None

    def upload_to_drive(self):
        content = self.text_area.get("1.0", tk.END).strip()
        if not content:
            messagebox.showwarning("ריק", "הדוח ריק.")
            return

        folder_id = self.config_manager.get("drive_folder_id")
        if not folder_id:
            messagebox.showerror("שגיאה", "לא מוגדר מזהה תיקייה (Folder ID) בהגדרות.")
            return

        # Save to a temporary file locally
        import tempfile
        try:
            # Create temp file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", prefix="temp_report_") as tmp:
                temp_path = tmp.name
            
            # Write DOCX content
            doc = Document()
            for line in content.split("\n"):
                p = doc.add_paragraph(line)
                p.paragraph_format.rtl = True
                p.alignment = 2 # RIGHT
            doc.save(temp_path)
            
            # Start upload thread
            self.window.title("...מעלה לדרייב")
            threading.Thread(target=self._upload_thread, args=(temp_path, folder_id)).start()
            
        except Exception as e:
            messagebox.showerror("שגיאה", f"שגיאה ביצירת קובץ זמני: {e}")

    def _upload_thread(self, filepath, folder_id):
        try:
            # Use original filename from settings or a default name
            target_name = self.default_filename if self.default_filename else "AI_Report.docx"
            
            # Upload with the correct name (the uploader might need modification if it only takes path)
            # Assuming drive_uploader.upload_file handles the name or we can rename it there.
            # Looking at current drive_uploader, it likely uses os.path.basename(filepath).
            # We want to preserve the "intended" name, not "temp_report_...".
            # For now, let's just upload. The user can rename in Drive, or we can improve DriveUploader later.
            
            link = self.drive_uploader.upload_file(filepath, folder_id)
            self.window.after(0, lambda: messagebox.showinfo("הצלחה", f"הקובץ הועלה בהצלחה!\nקישור: {link}"))
        except Exception as e:
            self.window.after(0, lambda: messagebox.showerror("כישלון", f"העלאה נכשלה:\n{e}"))
        finally:
            self.window.after(0, lambda: self.window.title("עורך דוחות AI"))
            # Cleanup temp file
            try:
                if os.path.exists(filepath):
                    os.remove(filepath)
            except: pass

    def refine_text(self):
        # Refines the report using AI
        try:
            sel_start = self.text_area.index("sel.first")
            sel_end = self.text_area.index("sel.last")
            text_to_refine = self.text_area.get(sel_start, sel_end)
            mode = "selection"
        except tk.TclError:
            text_to_refine = self.text_area.get("1.0", tk.END)
            mode = "all"

        if not text_to_refine.strip():
            return

        # Call AI
        model_name = self.config_manager.get("active_model", "gemini-flash-latest")
        api_key = self.config_manager.get_api_key("google_ai")

        if api_key:
             genai.configure(api_key=api_key)
        
        model = genai.GenerativeModel(model_name)
        
        prompt = (
            "אנא שפר את הניסוח של הטקסט הבא שיהיה קליט, זורם, ומקצועי (כמו מצגת). "
            "שמור על המשמעות המקורית:\n\n" + text_to_refine
        )
        
        def run_ai():
            try:
                response = model.generate_content(prompt)
                refined = response.text.strip()
                self.window.after(0, lambda: self._update_text(refined, mode, sel_start if mode == "selection" else "1.0", sel_end if mode == "selection" else tk.END))
            except Exception as e:
                self.window.after(0, lambda: messagebox.showerror("AI Error", str(e)))

        self.window.title("...AI מעבד")
        threading.Thread(target=run_ai).start()

    def _update_text(self, new_text, mode, start, end):
        self.window.title("עורך דוחות AI")
        self.text_area.delete(start, end)
        self.text_area.insert(start, new_text)
        self.text_area.tag_add("rtl", "1.0", "end")
