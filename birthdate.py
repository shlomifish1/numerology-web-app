# -*- coding: utf-8 -*-
import tkinter
from tkinter import Tk, Canvas, Entry, Button, PhotoImage, Radiobutton, StringVar, filedialog, messagebox, Toplevel
import personal_y  # Assuming this module contains This_Year class and its methods
import name  # Assuming this module contains NamesData class and its methods
from PIL import  ImageGrab
from tkinter import filedialog, messagebox

# --- DOCX Library Import ---
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # For text alignment

    DOCX_IMPORTED = True
except ImportError:
    print("Warning: 'python-docx' library not found. DOCX report generation will be disabled.")
    print("Install using: pip install python-docx")
    DOCX_IMPORTED = False
# -------------------------
import os
import sys
import traceback  # For better error details
import re  # For name validation
import subprocess  # For opening file/folder
from config_manager import ConfigManager
from settings_ui import SettingsUI
from chat_manager import ChatManager
from chat_ui import ChatUI
from report_editor import ReportEditor
from numerology_calculator import NumerologyCalculator


def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller. """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


class FullDates:

    def __init__(self):
        # --- Initialize all instance variables ---
        self.config_manager = ConfigManager()
        self.chat_manager = ChatManager(self.config_manager)
        self.calculator = NumerologyCalculator()
        self.fulldate_list_for_pyth = None
        self.real_year = None
        self.real_month = None
        self.full_list = None
        self.shana_ishit = None
        self.shana_nisteret = None  # Should store "X_Y" string
        self.bd_month = None
        self.first_pick = None;
        self.second_pick = None;
        self.third_pick = None;
        self.forth_pick = None
        self.first_challenge = None;
        self.second_challenge = None;
        self.third_challenge = None;
        self.forth_challenge = None
        self.peak1_reduced = None;
        self.peak2_reduced = None;
        self.peak3_reduced = None;
        self.peak4_reduced = None
        self.challenge1_reduced = None;
        self.challenge2_reduced = None;
        self.challenge3_reduced = None;
        self.challenge4_reduced = None
        self.first_pick_start = None;
        self.second_pick_start = None;
        self.third_pick_start = None;
        self.forth_pick_start = None
        self.final_number_destiny = None;
        self.original_destiny_sum = None
        self.p_year = None;
        self.p_month = None;
        self.p_day = None
        self.full_date_short = None
        self.window = None;
        self.canvas = None;
        self.entry_date = None;
        self.entry_first_name = None;
        self.entry_last_name = None
        self.gender_var = None;
        self.date_button = None;
        self.reload_button = None;
        self.save_options_button = None
        self.save_image_button = None;
        self.first_photo = None;
        self.second_photo = None;
        self.print_full_name_id = None
        self.fulldate_list = None;
        self.tzimtzum_age = None;
        self.age = None
        self.full_date = None;
        self.full_name = None;
        self.first_name_str = None;
        self.last_name_str = None
        self.first_name_val = None;
        self.full_name_val = None;
        self.itzurim_val = None;
        self.aiv_val = None
        self.first_quarter_reduced = None;
        self.second_quarter_reduced = None;
        self.third_quarter_reduced = None;
        self.forth_quarter_reduced = None
        self.bd_day_str = None;
        self.bd_month_str = None;
        self.bd_year_str = None



    # --- Helper methods moved to NumerologyCalculator ---
    # Wrappers kept for compatibility with existing code

    def get_interpretation(self, category, number, gender_folder_param, is_hidden_year=False,
                           is_peak_challenge_comb=False):
        return self.calculator.get_interpretation(category, number, gender_folder_param, is_hidden_year, is_peak_challenge_comb)

    def _set_rtl_paragraph(self, paragraph):
        """ Helper function to set paragraph to RTL and right-aligned. """
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.rtl = True

    def generate_docx_report(self):
        if not DOCX_IMPORTED:
            messagebox.showerror("שגיאה",
                                 "ספריית python-docx אינה מותקנת.\nלא ניתן ליצור דוח Word.\nאנא התקן באמצעות: pip install python-docx")
            return
        if self.p_day is None:
            messagebox.showerror("שגיאה", "אנא בצע חישוב מפה נומרולוגית תחילה לפני יצירת דוח.")
            return
        selected_gender = self.gender_var.get() if self.gender_var else "male"
        document = Document()

        heading = document.add_heading(f'דוח נומרולוגי עבור: {self.full_name}', level=0)
        self._set_rtl_paragraph(heading)

        p = document.add_paragraph(f"תאריך לידה: {self.entry_date.get()}")
        self._set_rtl_paragraph(p)
        p = document.add_paragraph(f"מין: {'זכר' if selected_gender == 'male' else 'נקבה'}")
        self._set_rtl_paragraph(p)
        p = document.add_paragraph(f"תאריך הפקת הדוח: {personal_y.today.strftime('%d/%m/%Y')}")
        self._set_rtl_paragraph(p)
        p = document.add_paragraph();
        self._set_rtl_paragraph(p)

        if self.final_number_destiny is not None:
            h = document.add_heading('מספר ייעוד (דרך חיים)', level=1);
            self._set_rtl_paragraph(h)
            p = document.add_paragraph(
                f"המספר המחושב: {self.final_number_destiny} (סכום ספרות מקורי: {self.full_date_short})");
            self._set_rtl_paragraph(p)
            interpretation = self.get_interpretation("destiny", self.final_number_destiny, selected_gender)
            p = document.add_paragraph(interpretation);
            self._set_rtl_paragraph(p);
            p = document.add_paragraph();
            self._set_rtl_paragraph(p)
        if self.p_day is not None:
            h = document.add_heading('יום לידה (לועזי)', level=1);
            self._set_rtl_paragraph(h)
            original_day = int(self.bd_day_str)
            p = document.add_paragraph(f"יום הלידה המקורי: {original_day}, לאחר צמצום: {self.p_day}");
            self._set_rtl_paragraph(p)
            interpretation_original = self.get_interpretation("birth_day", original_day, selected_gender)
            p = document.add_paragraph(f"פרשנות ליום הלידה {original_day}:\n{interpretation_original}");
            self._set_rtl_paragraph(p)
            if original_day != self.p_day:
                interpretation_reduced = self.get_interpretation("birth_day", self.p_day, selected_gender)
                p = document.add_paragraph(f"פרשנות ליום הלידה המצומצם {self.p_day}:\n{interpretation_reduced}");
                self._set_rtl_paragraph(p)
            p = document.add_paragraph();
            self._set_rtl_paragraph(p)
        if self.p_month is not None:
            h = document.add_heading('חודש לידה (לועזי)', level=1);
            self._set_rtl_paragraph(h)
            original_month = int(self.bd_month_str)
            p = document.add_paragraph(f"חודש הלידה המקורי: {original_month}, לאחר צמצום: {self.p_month}");
            self._set_rtl_paragraph(p)
            interpretation = self.get_interpretation("birth_month_reaction", original_month, selected_gender)
            p = document.add_paragraph(interpretation);
            self._set_rtl_paragraph(p);
            p = document.add_paragraph();
            self._set_rtl_paragraph(p)

        if self.first_name_val is not None:
            h = document.add_heading('מספר השם הפרטי', level=1);
            self._set_rtl_paragraph(h)
            p = document.add_paragraph(f"השם: {self.first_name_str}, הערך הנומרולוגי: {self.first_name_val}");
            self._set_rtl_paragraph(p)
            first_name_aspects = {
                "אפיון השם": "first_name_character",
                "נקודות לתיקון": "first_name_correcting",
                "אהבה וזוגיות": "first_name_love",
                "קריירה ועבודה": "first_name_work"
            }
            for sub_title, category_folder in first_name_aspects.items():
                interpretation_fn = self.get_interpretation(category_folder, self.first_name_val, selected_gender)
                if not interpretation_fn.startswith("[שגיאה") and not interpretation_fn.startswith(
                        "[קובץ פרשנות לא נמצא"):
                    h_sub = document.add_heading(sub_title, level=2);
                    self._set_rtl_paragraph(h_sub)
                    p_sub = document.add_paragraph(interpretation_fn);
                    self._set_rtl_paragraph(p_sub)
            p = document.add_paragraph();
            self._set_rtl_paragraph(p)

        if self.full_name_val is not None:
            h = document.add_heading('מספר השם המלא (שאיפת האדם)', level=1);
            self._set_rtl_paragraph(h)
            p = document.add_paragraph(f"השם המלא: {self.full_name}, הערך הנומרולוגי: {self.full_name_val}");
            self._set_rtl_paragraph(p)
            interpretation = self.get_interpretation("human_aspiration", self.full_name_val, selected_gender)
            p = document.add_paragraph(interpretation);
            self._set_rtl_paragraph(p);
            p = document.add_paragraph();
            self._set_rtl_paragraph(p)
        if self.itzurim_val is not None:
            h = document.add_heading('מספר העיצורים (האישיות החיצונית)', level=1);
            self._set_rtl_paragraph(h)
            p = document.add_paragraph(f"הערך הנומרולוגי: {self.itzurim_val}");
            self._set_rtl_paragraph(p)
            interpretation = self.get_interpretation("consonants", self.itzurim_val, selected_gender)
            p = document.add_paragraph(interpretation);
            self._set_rtl_paragraph(p);
            p = document.add_paragraph();
            self._set_rtl_paragraph(p)
        if self.aiv_val is not None:
            h = document.add_heading('מספר התנועות (הנשמה / הרצון הפנימי)', level=1);
            self._set_rtl_paragraph(h)
            p = document.add_paragraph(f"הערך הנומרולוגי: {self.aiv_val}");
            self._set_rtl_paragraph(p)
            interpretation = self.get_interpretation("vowels", self.aiv_val, selected_gender)
            p = document.add_paragraph(interpretation);
            self._set_rtl_paragraph(p);
            p = document.add_paragraph();
            self._set_rtl_paragraph(p)
        if self.shana_ishit is not None:
            h = document.add_heading('שנה אישית', level=1);
            self._set_rtl_paragraph(h)
            current_gregorian_year = personal_y.today.year
            p = document.add_paragraph(
                f"השנה האישית שלך לשנת {current_gregorian_year} (החל מיום ההולדת שלך בשנה זו): {self.shana_ishit}");
            self._set_rtl_paragraph(p)
            interpretation = self.get_interpretation("personal_year", self.shana_ishit, selected_gender)
            p = document.add_paragraph(interpretation);
            self._set_rtl_paragraph(p);
            p = document.add_paragraph();
            self._set_rtl_paragraph(p)
        if self.shana_nisteret is not None:
            h = document.add_heading('שנה נסתרת (פוטנציאל שנתי)', level=1);
            self._set_rtl_paragraph(h)
            display_shana_nisteret = str(self.shana_nisteret).replace('_', '/') if isinstance(self.shana_nisteret,
                                                                                              str) else self.shana_nisteret
            p = document.add_paragraph(f"השנה הנסתרת שלך: {display_shana_nisteret}");
            self._set_rtl_paragraph(p)
            interpretation = self.get_interpretation("hidden_year", str(self.shana_nisteret), selected_gender,
                                                     is_hidden_year=True)
            p = document.add_paragraph(interpretation);
            self._set_rtl_paragraph(p);
            p = document.add_paragraph();
            self._set_rtl_paragraph(p)

        if self.first_quarter_reduced is not None:
            h = document.add_heading('רבעונים אישיים לפי מיכל גרין', level=1)
            self._set_rtl_paragraph(h)

            quarters_data = [
                ("רבעון ראשון", self.first_quarter_reduced),
                ("רבעון שני", self.second_quarter_reduced),
                ("רבעון שלישי", self.third_quarter_reduced),
                ("רבעון רביעי", self.forth_quarter_reduced)
            ]

            for i, (label, quarter_val) in enumerate(quarters_data, 1):
                if quarter_val is None:
                    continue
                p = document.add_paragraph(f"{label}: {quarter_val}")
                self._set_rtl_paragraph(p)
                interpretation = self.get_interpretation("quarters", quarter_val, selected_gender)
                p = document.add_paragraph(interpretation)
                self._set_rtl_paragraph(p)

            p = document.add_paragraph()
            self._set_rtl_paragraph(p)

            # --- פירוש רצפים של 2 ספרות עם fallback הפיכה ---
            if hasattr(self, 'quarter_sequences_2'):
                p = document.add_heading('פירוש רצפים של 2 ספרות', level=2)
                self._set_rtl_paragraph(p)
                for original_seq in self.quarter_sequences_2:
                    # מנסים לקבל פירוש
                    interp = self.get_interpretation("quarters", original_seq, selected_gender)
                    used_seq = original_seq
                    # אם לא קיבלנו כלום, הופכים את הספרות
                    if not interp:
                        used_seq = original_seq[::-1]
                        interp = self.get_interpretation("quarters", used_seq, selected_gender)
                    # מדפיסים את הרצף (ואת הפירוש של הרצף שבחרנו)
                    p = document.add_paragraph(f"{original_seq} → {used_seq}:")
                    self._set_rtl_paragraph(p)
                    p = document.add_paragraph(interp or "אין פירוש לרצף זה")
                    self._set_rtl_paragraph(p)

            # --- פירוש רצפים של 3 ספרות עם אותו מנגנון fallback אם תרצו ---
            if hasattr(self, 'quarter_sequences_3'):
                p = document.add_heading('פירוש רצפים של 3 ספרות', level=2)
                self._set_rtl_paragraph(p)
                for original_seq in self.quarter_sequences_3:
                    interp = self.get_interpretation("quarters", original_seq, selected_gender)
                    used_seq = original_seq
                    if not interp:
                        used_seq = original_seq[::-1]
                        interp = self.get_interpretation("quarters", used_seq, selected_gender)
                    p = document.add_paragraph(f"{original_seq} → {used_seq}:")
                    self._set_rtl_paragraph(p)
                    p = document.add_paragraph(interp or "אין פירוש לרצף זה")
                    self._set_rtl_paragraph(p)

        h_main_peaks = document.add_heading('פסגות, אתגרים ושילובם', level=1);
        self._set_rtl_paragraph(h_main_peaks)
        peaks_challenges_combined_data = [
            ("תקופה ראשונה", self.peak1_reduced, self.challenge1_reduced, f"מגיל לידה עד גיל {self.first_pick_start}"),
            ("תקופה שנייה", self.peak2_reduced, self.challenge2_reduced,
             f"מגיל {self.first_pick_start + 1} עד גיל {self.second_pick_start - 1}"),
            ("תקופה שלישית", self.peak3_reduced, self.challenge3_reduced,
             f"מגיל {self.second_pick_start} עד גיל {self.third_pick_start - 1}"),
            ("תקופה רביעית", self.peak4_reduced, self.challenge4_reduced, f"מגיל {self.third_pick_start} ואילך")
        ]
        for period_name, peak_num, challenge_num, age_range in peaks_challenges_combined_data:
            h_period = document.add_heading(f"{period_name} ({age_range})", level=2);
            self._set_rtl_paragraph(h_period)
            if peak_num is not None:
                p = document.add_paragraph(f"פסגה: {peak_num}");
                self._set_rtl_paragraph(p)
                interpretation_peak = self.get_interpretation("peaks_interpretation", peak_num, selected_gender)
                p = document.add_paragraph(interpretation_peak);
                self._set_rtl_paragraph(p)
            if challenge_num is not None:
                p = document.add_paragraph(f"אתגר: {challenge_num}");
                self._set_rtl_paragraph(p)
                interpretation_challenge = self.get_interpretation("challenges_interpretation", challenge_num,
                                                                   selected_gender)
                p = document.add_paragraph(interpretation_challenge);
                self._set_rtl_paragraph(p)
            if peak_num is not None and challenge_num is not None:
                combined_val_str = f"{peak_num}{challenge_num}"
                p = document.add_paragraph(f"שילוב פסגה-אתגר ({combined_val_str}):");
                self._set_rtl_paragraph(p)
                interpretation_combined = self.get_interpretation("peak_challenge_comb", combined_val_str,
                                                                  selected_gender, is_peak_challenge_comb=True)
                p = document.add_paragraph(interpretation_combined);
                self._set_rtl_paragraph(p)
            p = document.add_paragraph();
            self._set_rtl_paragraph(p)

        h_pytha = document.add_heading('ריבוע פיתגורס', level=1);
        self._set_rtl_paragraph(h_pytha)
        birth_date_str_for_pytha = self.entry_date.get()
        pytha_interpretation_text = ""
        if len(birth_date_str_for_pytha) == 8 and birth_date_str_for_pytha.isdigit():
            intro_pytha = self.get_interpretation("pythagorean_square", "intro", selected_gender)
            if not intro_pytha.startswith("["): pytha_interpretation_text += intro_pytha + "\n\n"
            digit_counts = {}
            for digit_char in birth_date_str_for_pytha:
                digit_int = int(digit_char)
                if digit_int != 0: digit_counts[digit_int] = digit_counts.get(digit_int, 0) + 1
            pytha_interpretation_text += "המספרים המופיעים בתאריך הלידה (ולכן בריבוע):\n"
            for i in range(1, 10):
                if i in digit_counts:
                    pytha_interpretation_text += f"ספרה {i} מופיעה {digit_counts[i]} פעמים. "
                    interpretation_digit = self.get_interpretation("pythagorean_square", f"number_{i}_meaning",
                                                                   selected_gender)
                    if not interpretation_digit.startswith("["): pytha_interpretation_text += interpretation_digit + " "
                else:
                    pytha_interpretation_text += f"ספרה {i} חסרה. "
                if i % 3 == 0: pytha_interpretation_text += "\n"
            pytha_interpretation_text += "\n"
        else:
            pytha_interpretation_text += "לא ניתן היה לעבד את ריבוע פיתגורס עקב תאריך לא תקין."
        p = document.add_paragraph(pytha_interpretation_text);
        self._set_rtl_paragraph(p);
        p = document.add_paragraph();
        self._set_rtl_paragraph(p)
        try:
            default_filename = f"דוח נומרולוגי - {self.full_name} - {self.entry_date.get()}.docx"
            default_filename = "".join(
                c if c.isalnum() or c in (' ', '-', '_', '.', '(', ')', '[', ']') else '_' for c in default_filename)
            filepath = filedialog.asksaveasfilename(initialfile=default_filename, defaultextension=".docx",
                                                    filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")])
            if filepath:
                document.save(filepath)
                success_message = f"הדוח נשמר בהצלחה במיקום:\n{filepath}"
                dialog = tkinter.Toplevel(self.window);
                dialog.title("הדוח נשמר");
                dialog.geometry("450x150");
                dialog.resizable(False, False)
                window_x = self.window.winfo_x();
                window_y = self.window.winfo_y();
                window_width = self.window.winfo_width();
                window_height = self.window.winfo_height()
                dialog_x = window_x + (window_width // 2) - (450 // 2);
                dialog_y = window_y + (window_height // 2) - (150 // 2);
                dialog.geometry(f"+{dialog_x}+{dialog_y}")
                label = tkinter.Label(dialog, text=success_message, wraplength=400, justify="center", pady=10);
                label.pack()
                button_frame = tkinter.Frame(dialog);
                button_frame.pack(pady=10)

                def open_file_action():
                    try:
                        if sys.platform == "win32":
                            os.startfile(filepath)
                        elif sys.platform == "darwin":
                            subprocess.run(['open', filepath])
                        else:
                            subprocess.run(['xdg-open', filepath])
                    except Exception as e_open:
                        messagebox.showerror("שגיאה בפתיחת קובץ", f"לא ניתן היה לפתוח את הקובץ:\n{e_open}",
                                             parent=dialog)
                    dialog.destroy()

                def open_folder_action():
                    try:
                        if sys.platform == "win32":
                            subprocess.run(['explorer', '/select,', os.path.normpath(filepath)])
                        elif sys.platform == "darwin":
                            subprocess.run(['open', '-R', filepath])
                        else:
                            subprocess.run(['xdg-open', os.path.dirname(filepath)])
                    except Exception as e_folder:
                        messagebox.showerror("שגיאה בפתיחת תיקייה", f"לא ניתן היה לפתוח את התיקייה:\n{e_folder}",
                                             parent=dialog)
                    dialog.destroy()

                def close_dialog_action():
                    dialog.destroy()

                btn_open_file = tkinter.Button(button_frame, text="פתח את הקובץ", command=open_file_action, width=15);
                btn_open_file.pack(side="left", padx=5)
                btn_open_folder = tkinter.Button(button_frame, text="פתח תיקייה", command=open_folder_action, width=15);
                btn_open_folder.pack(side="left", padx=5)
                btn_ok = tkinter.Button(button_frame, text="אישור", command=close_dialog_action, width=10);
                btn_ok.pack(side="left", padx=5)
                dialog.grab_set();
                dialog.wait_window()
            else:
                messagebox.showinfo("בוטל", "שמירת הדוח בוטלה.")
        except Exception as e:
            error_msg = f"שגיאה בשמירת קובץ ה-DOCX:\n{e}\n{traceback.format_exc()}"
            print(f"Error saving DOCX: {error_msg}");
            messagebox.showerror("שגיאת שמירה", error_msg)

    def _show_post_save_popup(self, filepath):
        # יוצרים דיאלוג
        dialog = tkinter.Toplevel(self.window)
        dialog.title("הדוח נשמר")
        dialog.geometry("450x150")
        dialog.resizable(False, False)
        # ממרכזים אותו מעל החלון הראשי
        window_x = self.window.winfo_x()
        window_y = self.window.winfo_y()
        window_w = self.window.winfo_width()
        window_h = self.window.winfo_height()
        dialog_x = window_x + (window_w  // 2) - (450 // 2)
        dialog_y = window_y + (window_h // 2) - (150 // 2)
        dialog.geometry(f"+{dialog_x}+{dialog_y}")

        # כותרת טקסט
        label = tkinter.Label(
            dialog,
            text=f"הדוח נשמר בהצלחה במיקום:\n{filepath}",
            wraplength=400,
            justify="center",
            pady=10
        )
        label.pack()

        # מסגרת לכפתורים
        button_frame = tkinter.Frame(dialog)
        button_frame.pack(pady=10)

        # callbacks של הכפתורים
        def open_file_action():
            try:
                if sys.platform == "win32":
                    os.startfile(filepath)
                elif sys.platform == "darwin":
                    subprocess.run(["open", filepath])
                else:
                    subprocess.run(["xdg-open", filepath])
            except Exception as e_open:
                messagebox.showerror("שגיאה בפתיחת קובץ", f"לא ניתן לפתוח את הקובץ:\n{e_open}", parent=dialog)
            dialog.destroy()

        def open_folder_action():
            try:
                folder = os.path.dirname(filepath)
                if sys.platform == "win32":
                    os.startfile(folder)
                elif sys.platform == "darwin":
                    subprocess.run(["open", "-R", filepath])
                else:
                    subprocess.run(["xdg-open", folder])
            except Exception as e_folder:
                messagebox.showerror("שגיאה בפתיחת תיקייה", f"לא ניתן לפתוח את התיקייה:\n{e_folder}", parent=dialog)
            dialog.destroy()

        def close_dialog_action():
            dialog.destroy()

        # יצירת הכפתורים
        btn_open_file = tkinter.Button(
            button_frame, text="פתח את הקובץ",
            command=open_file_action, width=15
        )
        btn_open_file.pack(side="left", padx=5)

        btn_open_folder = tkinter.Button(
            button_frame, text="פתח תיקייה",
            command=open_folder_action, width=15
        )
        btn_open_folder.pack(side="left", padx=5)

        btn_ok = tkinter.Button(
            button_frame, text="אישור",
            command=close_dialog_action, width=10
        )
        btn_ok.pack(side="left", padx=5)

        # תפיסת הפוקוס עד לסגירת הדיאלוג
        dialog.grab_set()
        self.window.wait_window(dialog)


    def generate_gpt_docx_report(self):
        if not DOCX_IMPORTED:
            messagebox.showerror(
                "שגיאה",
                "ספריית python-docx אינה מותקנת. התקן: pip install python-docx"
            )
            return

        # 1) לבנות את המילון כמו בשיטה הקיימת
        data = {
            "full_name":           self.full_name,
            "birth_date":          self.full_date,
            "personal_day":        self.p_day,
            "personal_month":      self.p_month,
            "personal_year":       self.p_year,
            "destiny_number":      self.final_number_destiny,
            "personal_year_number":self.shana_ishit,
            "hidden_year":         self.shana_nisteret,
            "age":                 self.age,
            "life_peaks": [
                self.peak1_reduced, self.peak2_reduced,
                self.peak3_reduced, self.peak4_reduced
            ],
            "challenges": [
                self.challenge1_reduced, self.challenge2_reduced,
                self.challenge3_reduced, self.challenge4_reduced
            ],
            "quarters": [
                self.first_quarter_reduced, self.second_quarter_reduced,
                self.third_quarter_reduced, self.forth_quarter_reduced
            ],
        }

        # 2) לקרוא ל־AI
        from gpt_report import generate_person_report
        
        # Load from config
        active_model = self.config_manager.get("active_model", "gemini-flash-latest")
        google_key = self.config_manager.get_api_key("google_ai")

        report_text = generate_person_report(data, model_name=active_model, api_key=google_key)

        # 3) פתח עורך דוחות (במקום שמירה ישירה)
        default_filename = f"דוח AI - {self.full_name} - {self.entry_date.get()}.docx"
        default_filename = "".join(c for c in default_filename if c not in '<>:"/\\|?*') # Sanitize
        
        ReportEditor(self.window, report_text, self.config_manager, default_filename)

    def show_save_options(self):
        """ This method is called by the 'צור דוח' button. """
        print("Save options button clicked - attempting DOCX generation...")
        self.generate_docx_report()

    def calc_quarters_michal_green(self):
        try:
            # ודא שהנתונים הדרושים קיימים
            if self.final_number_destiny is None or self.shana_ishit is None:
                print("שנה אישית או מספר גורל לא מחושבים – לא ניתן לחשב רבעונים.")
                return

            # חישוב השנה האוניברסלית
            current_year = personal_y.today.year
            universal_year = sum(int(d) for d in str(current_year))  # למשל 2+0+2+5 = 9

            # חישוב ראשוני של כל אחד מהרבעונים
            q1 = self.final_number_destiny + universal_year  # גורל + אוניברסלי
            q2 = self.final_number_destiny + self.shana_ishit  # גורל + שנה אישית
            q3 = q1 + q2  # רבעון ראשון + שני
            q4 = self.shana_ishit + universal_year  # שנה אישית + אוניברסלי

            # צמצום כל תוצאה לספרה בודדת
            self.first_quarter_reduced = self.short_number_single(q1)
            self.second_quarter_reduced = self.short_number_single(q2)
            self.third_quarter_reduced = self.short_number_single(q3)
            self.forth_quarter_reduced = self.short_number_single(q4)

            # בניית הרצפים
            quarters = [
                self.first_quarter_reduced,
                self.second_quarter_reduced,
                self.third_quarter_reduced,
                self.forth_quarter_reduced
            ]
            # רצפים של 2 ספרות (חלון נע מחלק 2)
            self.quarter_sequences_2 = [
                f"{quarters[i]}{quarters[i + 1]}"
                for i in range(len(quarters) - 1)
            ]
            # רצפים של 3 ספרות (חלון נע של 3)
            self.quarter_sequences_3 = [
                f"{quarters[i]}{quarters[i + 1]}{quarters[i + 2]}"
                for i in range(len(quarters) - 2)
            ]

        except Exception as e:
            print(f"שגיאה בחישוב רבעונים לפי מיכל גרין: {e}")

    def bd_split(self, event=None):
        self.full_date = self.entry_date.get()
        self.first_name_str = self.entry_first_name.get().strip()
        self.last_name_str = self.entry_last_name.get().strip()
        selected_gender = self.gender_var.get() if self.gender_var else None

        # --- Validation ---
        if not (len(self.full_date) == 8 and self.full_date.isdigit()):
            messagebox.showerror("שגיאת קלט", "פורמט תאריך לא תקין.\nאנא הזן 8 ספרות ללא רווחים (לדוגמה: 25081988).")
            return
        if not self.first_name_str: 
            messagebox.showerror("שגיאת קלט", "שם פרטי הוא שדה חובה.")
            return
        if not re.match(r"^[a-zA-Z\u0590-\u05FF\s]+$", self.first_name_str): 
            messagebox.showerror("שגיאת קלט", "שם פרטי יכול להכיל רק אותיות ורווחים.")
            return
        if not self.last_name_str: 
            messagebox.showerror("שגיאת קלט", "שם משפחה הוא שדה חובה.")
            return
        if not re.match(r"^[a-zA-Z\u0590-\u05FF\s]+$", self.last_name_str): 
            messagebox.showerror("שגיאת קלט", "שם משפחה יכול להכיל רק אותיות ורווחים.")
            return
        if not selected_gender: 
            messagebox.showerror("שגיאת קלט", "אנא בחר מין (זכר/נקבה).")
            return

        try:
            # --- Use Calculator ---
            day_str = self.full_date[0:2]
            month_str = self.full_date[2:4]
            year_str = self.full_date[4:]
            
            # Calculate EVERYTHING
            self.calculator.calculate(day_str, month_str, year_str, self.first_name_str, self.last_name_str, selected_gender)
            
            # --- Sync State Back to UI Class ---
            c = self.calculator
            self.p_day = c.p_day
            self.p_month = c.p_month
            self.p_year = c.p_year
            self.final_number_destiny = c.final_number_destiny
            self.full_date_short = c.full_date_short
            self.full_name = c.full_name
            self.age = c.age
            self.shana_ishit = c.shana_ishit
            self.shana_nisteret = c.shana_nisteret
            
            self.first_name_val = c.first_name_val
            self.full_name_val = c.full_name_val
            self.itzurim_val = c.itzurim_val
            self.aiv_val = c.aiv_val
            self.tzimtzum_age = c.tzimtzum_age
            
            # Peaks
            self.peak1_reduced = c.peak1_reduced
            self.peak2_reduced = c.peak2_reduced
            self.peak3_reduced = c.peak3_reduced
            self.peak4_reduced = c.peak4_reduced
            self.first_pick_start = c.first_pick_start
            self.second_pick_start = c.second_pick_start
            self.third_pick_start = c.third_pick_start
            self.forth_pick_start = c.forth_pick_start
            
            # Challenges
            self.challenge1_reduced = c.challenge1_reduced
            self.challenge2_reduced = c.challenge2_reduced
            self.challenge3_reduced = c.challenge3_reduced
            self.challenge4_reduced = c.challenge4_reduced
            
            # Quarters
            self.first_quarter_reduced = c.first_quarter_reduced
            self.second_quarter_reduced = c.second_quarter_reduced
            self.third_quarter_reduced = c.third_quarter_reduced
            self.forth_quarter_reduced = c.forth_quarter_reduced
            self.quarter_sequences_2 = c.quarter_sequences_2
            self.quarter_sequences_3 = c.quarter_sequences_3
            
            # Raw string values for DOCX use if needed (legacy)
            self.bd_day_str = day_str
            self.bd_month_str = month_str
            self.bd_year_str = year_str
            
            # --- Update AI Context ---
            self.build_ai_context()
            
            # --- Display Results ---
            self.print_final_details()

        except Exception as e:
            traceback.print_exc()
            messagebox.showerror("שגיאת חישוב", f"אירעה שגיאה בחישוב:\n{e}")

    def build_ai_context(self):
        """ Gathers relevant interpretations and feeds them to the Chat. """
        try:
            context_parts = []
            gender = self.gender_var.get()
            
            # 1. Personal Year
            if self.shana_ishit:
                text = self.calculator.get_interpretation("personal_year", self.shana_ishit, gender)
                if text and not text.startswith("["):
                    context_parts.append(f"--- שנה אישית {self.shana_ishit} ---\n{text}")

            # 2. Destiny
            if self.final_number_destiny:
                text = self.calculator.get_interpretation("destiny", self.final_number_destiny, gender)
                if text and not text.startswith("["):
                    context_parts.append(f"--- מספר ייעוד {self.final_number_destiny} ---\n{text}")

            # 3. Peak 1 (Mental/Example)
            if self.peak1_reduced:
                text = self.calculator.get_interpretation("peaks_interpretation", self.peak1_reduced, gender)
                if text and not text.startswith("["):
                     context_parts.append(f"--- פסגה ראשונה {self.peak1_reduced} ---\n{text}")

            full_context = f"נתוני המשתמש הנוכחי (Context Profile):\nשם: {self.full_name}\nתאריך לידה: {self.full_date}\n\n"
            full_context += "\n\n".join(context_parts)
            
            self.chat_manager.set_context(full_context)
            print("AI Context Updated.")
        except Exception as e:
            print(f"Error building AI context: {e}")




    def print_final_details(self):
        # בדיקה ראשונית - אם אין יום לידה, מציגים מסך שגיאה
        if self.p_day is None:
            self.canvas.delete("all")
            try:
                self.second_photo = PhotoImage(file=resource_path("image/second_page.png"))
                self.canvas.create_image(0, 0, image=self.second_photo, anchor="nw")
            except Exception: pass
            self.canvas.create_text(500, 300, text="אנא בצע חישוב תחילה.", font=("Arial", 16, "bold"), fill="red")
            return
        
        # ניקוי וסידור ה-Canvas (ללא PanedWindow)
        self.canvas.delete("all")
        if hasattr(self, 'paned_window'):
            self.paned_window.destroy()
            del self.paned_window
        
        self.canvas.pack(fill=tkinter.BOTH, expand=True)
        
        # --- טעינת תמונת הרקע ---
        try:
            self.second_photo = PhotoImage(file=resource_path("image/second_page.png"))
            self.canvas.create_image(0, 0, image=self.second_photo, anchor="nw")
        except Exception as e:
            print(f"Error loading second_page.png: {e}")
            self.canvas.create_text(500, 300, text="Error loading background image.", font=("Arial", 16), fill="red")

        # --- הצגת נתונים ---
        try:
            year_code_display = f"{self.age}{self.tzimtzum_age}{self.shana_ishit}" if self.age else "N/A"
            Y_OFFSET = 0
            FONT_INFO = ("Arial", 20, "italic")
            FONT_TITLE = ("Arial", 25, "bold italic")
            FILL_COLOR = "white"
            TITLE_FILL = "yellow"

            self.print_full_name_id = self.canvas.create_text(800, 40, fill=TITLE_FILL, text=self.full_name, font=FONT_TITLE, anchor="center")
            self.canvas.create_text(580, 40, fill=TITLE_FILL, text=self.entry_date.get(), font=FONT_TITLE, anchor="center")
            
            date_details = f"{self.p_day} / {self.p_month} / {self.p_year}"
            self.canvas.create_text(580, 100, fill=FILL_COLOR, text=date_details, font=FONT_INFO, anchor="center")
            
            destiny = f"{self.final_number_destiny} / {self.full_date_short}"
            self.canvas.create_text(790, 150, fill=FILL_COLOR, text=destiny, font=FONT_INFO, anchor="center")

            self.canvas.create_text(800, 209, fill=FILL_COLOR, text=str(self.first_name_val), font=FONT_INFO)
            self.canvas.create_text(635, 265, fill=FILL_COLOR, text=str(self.full_name_val), font=FONT_INFO)
            self.canvas.create_text(825, 320, fill=FILL_COLOR, text=str(self.itzurim_val), font=FONT_INFO)
            self.canvas.create_text(870, 375, fill=FILL_COLOR, text=str(self.aiv_val), font=FONT_INFO)
            self.canvas.create_text(790, 425, fill=FILL_COLOR, text=str(self.shana_ishit), font=FONT_INFO)
            
            nisteret = f"{self.shana_nisteret.split('_')[1]}/{self.final_number_destiny}" if self.shana_nisteret else "N/A"
            self.canvas.create_text(790, 479, fill=FILL_COLOR, text=nisteret, font=FONT_INFO)
            self.canvas.create_text(795, 525, fill=FILL_COLOR, text=str(self.age), font=FONT_INFO)
            self.canvas.create_text(795, 565, fill=FILL_COLOR, text=year_code_display, font=FONT_INFO)

            # Peaks (Left side)
            self.canvas.create_text(90, 135, fill=FILL_COLOR, text=str(self.peak1_reduced), font=FONT_INFO)
            self.canvas.create_text(90, 185, fill=FILL_COLOR, text=str(self.peak2_reduced), font=FONT_INFO)
            self.canvas.create_text(90, 235, fill=FILL_COLOR, text=str(self.peak3_reduced), font=FONT_INFO)
            self.canvas.create_text(90, 285, fill=FILL_COLOR, text=str(self.peak4_reduced), font=FONT_INFO)

            self.canvas.create_text(230, 135, fill=FILL_COLOR, text=str(self.challenge1_reduced), font=FONT_INFO)
            self.canvas.create_text(230, 185, fill=FILL_COLOR, text=str(self.challenge2_reduced), font=FONT_INFO)
            self.canvas.create_text(230, 235, fill=FILL_COLOR, text=str(self.challenge3_reduced), font=FONT_INFO)
            self.canvas.create_text(230, 285, fill=FILL_COLOR, text=str(self.challenge4_reduced), font=FONT_INFO)

            # Ages
            self.canvas.create_text(370, 135, fill=FILL_COLOR, text=f"{self.first_pick_start} - {self.second_pick_start - 1}", font=FONT_INFO, anchor="center")
            self.canvas.create_text(370, 185, fill=FILL_COLOR, text=f"{self.second_pick_start} - {self.third_pick_start - 1}", font=FONT_INFO, anchor="center")
            self.canvas.create_text(370, 235, fill=FILL_COLOR, text=f"{self.third_pick_start} - {self.forth_pick_start - 1}", font=FONT_INFO, anchor="center")
            self.canvas.create_text(370, 285, fill=FILL_COLOR, text=f"{self.forth_pick_start} - {self.forth_pick_start+8}", font=FONT_INFO, anchor="center")

            # Quarters
            self.canvas.create_text(370, 375, fill=FILL_COLOR, text=str(self.first_quarter_reduced), font=FONT_INFO)
            self.canvas.create_text(370, 420, fill=FILL_COLOR, text=str(self.second_quarter_reduced), font=FONT_INFO)
            self.canvas.create_text(370, 465, fill=FILL_COLOR, text=str(self.third_quarter_reduced), font=FONT_INFO)
            self.canvas.create_text(370, 505, fill=FILL_COLOR, text=str(self.forth_quarter_reduced), font=FONT_INFO)

            # Pythagoras
            def get_x_pytha(n): return {1: 530, 2: 540, 3: 542, 4: 625, 5: 626, 6: 627, 7: 705, 8: 706, 9: 707}.get(n)
            def get_y_pytha(n): return {1: 375, 4: 376, 7: 378, 2: 420, 5: 420, 8: 420, 3: 463, 6: 460, 9: 464}.get(n)
            
            bd = self.entry_date.get()
            if len(bd) == 8 and bd.isdigit():
                counts = {}
                for c in bd: counts[int(c)] = counts.get(int(c), 0) + 1
                for d, count in counts.items():
                    if d!=0: self.canvas.create_text(get_x_pytha(d), get_y_pytha(d), fill=FILL_COLOR, text=str(d)*count, font=FONT_INFO, anchor="center")

            # --- כפתורים למטה ---
            BTN_FONT = ('Arial', 12, 'bold'); BTN_BG='#1a1a1a'; BTN_FG='white'; BTN_W=13
            buttons_y = 650; spacing = 120; center_x = self.canvas.winfo_width() / 2
            if center_x < 100: center_x = 500 # fallback if not fully drawn yet

            # 1. שמור תמונה
            self.save_image_button = Button(self.window, text="שמור תמונה", command=self.save_map_image, width=BTN_W, bg=BTN_BG, fg=BTN_FG, font=BTN_FONT)
            self.canvas.create_window(center_x - 2*spacing, buttons_y, window=self.save_image_button)
            
            # 2. צור דוח רגיל
            self.save_options_button = Button(self.window, text="צור דוח", command=self.show_save_options, width=BTN_W, bg=BTN_BG, fg=BTN_FG, font=BTN_FONT)
            self.canvas.create_window(center_x - 1*spacing, buttons_y, window=self.save_options_button)
            
            # 3. צור דוח AI
            self.save_gpt_button = Button(self.window, text="AI צור דוח", command=self.generate_gpt_docx_report, width=BTN_W, bg=BTN_BG, fg=BTN_FG, font=BTN_FONT)
            self.canvas.create_window(center_x, buttons_y, window=self.save_gpt_button)

            # 4. כפתור צ'אט (חדש)
            self.chat_button = Button(self.window, text="צ'אט AI", command=self.open_chat_window, width=BTN_W, bg=BTN_BG, fg=BTN_FG, font=BTN_FONT)
            self.canvas.create_window(center_x + 1*spacing, buttons_y, window=self.chat_button)
            
            # 5. הפעל מחדש
            self.reload_button = Button(self.window, text="הפעל מחדש", command=self.restart_program, width=BTN_W, bg=BTN_BG, fg=BTN_FG, font=BTN_FONT)
            self.canvas.create_window(center_x + 2*spacing, buttons_y, window=self.reload_button)
            
            # 6. כפתור הגדרות (מיקום בפינה שמאלית עליונה כדי לא להסתיר)
            self.settings_button = Button(self.window, text="הגדרות", command=lambda: SettingsUI(self.window, self.config_manager), width=8, bg=BTN_BG, fg=BTN_FG, font=('Arial', 10))
            self.canvas.create_window(50, 55, window=self.settings_button) # Top-Left corner, moved down
            
            self.window.bind('<F5>', lambda e: self.restart_program())

        except Exception as e:
            print(f"Error displaying details: {e}")
            messagebox.showerror("שגיאה", f"תקלה בתצוגה:\n{e}")

    def open_chat_window(self):
        """ Opens the AI Chat in a separate top-level window. """
        if hasattr(self, 'chat_window') and self.chat_window.winfo_exists():
            self.chat_window.lift()
            return
        
        self.chat_window = Toplevel(self.window)
        self.chat_window.title("AI Chat - Numerology Assistant")
        self.chat_window.geometry("400x600")
        
        chat_ui = ChatUI(self.chat_window, self.chat_manager)
        chat_ui.pack(fill=tkinter.BOTH, expand=True)

    def get_date(self):
        self.window = Tk()
        self.window.geometry("1000x700")  # width x height in pixels
        self.window.title("מחשבון נומרולוגיה - פיתוח והרחבה")
        self.window.minsize(1000, 700)
        self.gender_var = StringVar(self.window, value="male")
        try:
            self.first_photo = PhotoImage(file=resource_path("image/first_page.png"))
            iw = self.first_photo.width();
            ih = self.first_photo.height()
            self.canvas = Canvas(self.window, width=iw, height=ih, bd=0, highlightthickness=0, bg="black")
            self.canvas.create_image(0, 0, image=self.first_photo, anchor="nw");
            self.canvas.pack(fill="both", expand=True)
        except Exception as e:
            print(f"Error loading first_page.png: {e}")
            iw, ih = 1000, 700;
            self.canvas = Canvas(self.window, width=iw, height=ih, bg="#2B2B2B");
            self.canvas.pack(fill="both", expand=True)
            self.canvas.create_text(iw / 2, 50, text="מחשבון נומרולוגיה", font=("Arial", 30, "bold"), fill="white",
                                    anchor="center")
            self.canvas.create_text(iw / 2, ih / 2 - 50, text="שגיאה בטעינת תמונת רקע.\nהתוכנה עדיין פעילה.",
                                    fill="orange", font=("Arial", 16), anchor="center", justify="center")
        # הגדרות ראשוניות (כפי שסיפקת)
        ex = 850
        lxo = 130
        eys = 150
        eyst = 85
        ef = ('Arial', 16)
        ew = 20
        lf = ('Arial', 14, 'bold')
        lfc = "white"

        # חישובי מיקומים חדשים
        x_label = ex  # תווית תתחיל כאן (anchor="w")
        x_entry = ex - lxo  # שדה הקלט יסתיים כאן (anchor="e")
        start_y = eys + 120  # מעלה/מטה למרכז האנכי
        row_gap = eyst  # רווח בין שורות

        # ——— שורה 1: תאריך לידה ———
        # Entry
        dvv = StringVar()

        def vdicbc(P):
            return False if len(P) > 8 else (P.isdigit() or P == "")

        vdc = (self.window.register(vdicbc), '%P')
        self.entry_date = Entry(
            self.window, font=ef, textvariable=dvv,
            width=ew, justify='center', bd=2,
            validate='key', validatecommand=vdc
        )
        self.canvas.create_window(
            x_entry, start_y,
            window=self.entry_date,
            anchor="e"
        )

        # תווית
        self.canvas.create_text(
            x_label, start_y,
            text=":תאריך לידה", font=lf,
            fill=lfc, anchor="w"
        )

        # Placeholder
        self.canvas.create_text(
            x_entry - 10, start_y,
            text="(DDMMYYYY)", font=('Arial', 10, 'italic'),
            fill=lfc, anchor="e"
        )

        # ——— שורה 2: שם פרטי ———
        fnv = StringVar()

        def vnicbc(P):
            return True if P == "" else bool(re.fullmatch(r"^[a-zA-Z\u0590-\u05FF\s]*$", P))

        vnc = (self.window.register(vnicbc), '%P')
        self.entry_first_name = Entry(
            self.window, font=ef, textvariable=fnv,
            width=ew, justify='center', bd=2,
            validate='key', validatecommand=vnc
        )
        self.canvas.create_window(
            x_entry, start_y + row_gap,
            window=self.entry_first_name,
            anchor="e"
        )
        self.canvas.create_text(
            x_label, start_y + row_gap,
            text=":שם פרטי", font=lf,
            fill=lfc, anchor="w"
        )

        # ——— שורה 3: שם משפחה ———
        lnv = StringVar()
        self.entry_last_name = Entry(
            self.window, font=ef, textvariable=lnv,
            width=ew, justify='center', bd=2,
            validate='key', validatecommand=vnc
        )
        self.canvas.create_window(
            x_entry, start_y + 2 * row_gap,
            window=self.entry_last_name,
            anchor="e"
        )
        self.canvas.create_text(
            x_label, start_y + 2 * row_gap,
            text=":שם משפחה", font=lf,
            fill=lfc, anchor="w"
        )

        # קבע את מיקום תחילת הטופס (יורד קצת למטה)
        start_y = eys + 120
        row_gap = eyst

        # ——— שורה 4: "מין:" + Radiobuttons ———
        gyp = start_y + 3 * row_gap
        # תווית "מין:"
        self.canvas.create_text(
            x_label +30, start_y + 3 * row_gap,
            text=":מין", font=lf, fill=lfc,
            anchor="e"
        )

        # הגדרות הרדיו־בוטונים
        rbf = ('Arial', 12, 'bold')
        rbbg = "#4a4a4a"
        rbfg = "white"
        rbsc = "#606060"

        mrb = Radiobutton(
            self.window, text="זכר",
            variable=self.gender_var, value="male",
            font=rbf, bg=rbbg, fg=rbfg,
            selectcolor=rbsc, activebackground=rbsc,
            activeforeground=rbfg, indicatoron=0,
            relief="flat", padx=10, pady=5
        )
        frb = Radiobutton(
            self.window, text="נקבה",
            variable=self.gender_var, value="female",
            font=rbf, bg=rbbg, fg=rbfg,
            selectcolor=rbsc, activebackground=rbsc,
            activeforeground=rbfg, indicatoron=0,
            relief="flat", padx=10, pady=5
        )

        # מיקום הבוטונים משמאל ל־ex ו־80 פיקסל ימינה ממנו
        self.canvas.create_window(ex -250, gyp, window=mrb, anchor="w")
        self.canvas.create_window(ex - 350, gyp, window=frb, anchor="w")

        # ——— שורה 5: כפתור "חשב מפה נומרולוגית" ———
        byp = gyp + row_gap + 10  # עוד שורה מטה + מרחק של 10 פיקסל

        self.date_button = Button(
            self.window,
            command=self.bd_split,
            text="חשב מפה נומרולוגית",
            width=20, height=1,
            bd=3, relief='raised',
            bg='#2a5298', fg='white',
            font=('Arial', 16, 'bold'),
            activebackground='#3b6ec2',
            activeforeground='white'
        )
        # מרכז האופקי – מספיק להשתמש ב־ex
        self.canvas.create_window(
            ex -250, byp,
            window=self.date_button,
            anchor="center"
        )

        self.window.bind('<Return>', self.bd_split);
        self.entry_date.focus_set();
        self.window.mainloop()

    def restart_program(self):
        if self.window: self.window.destroy()
        self.__init__();
        self.get_date()

    def save_map_image(self):
        if self.p_day is None:
            messagebox.showerror("שגיאה", "לא ניתן לשמור תמונה. אנא בצע חישוב מפה תחילה.")
            return
        if not hasattr(self, 'print_full_name_id') or not self.print_full_name_id:
            messagebox.showerror("שגיאה", "לא ניתן לשמור תמונה, נתוני שם לא אותרו על הקנבס.")
            return
        try:
            item_text = self.canvas.itemcget(self.print_full_name_id, 'text')
            if not item_text:
                print("Warning: Full name item exists but has no text. Using default filename for image.")
        except tkinter.TclError:
            messagebox.showerror("שגיאה", "לא ניתן לשמור תמונה, פריט השם המלא לא נמצא על הקנבס.")
            return
        except Exception as e:
            messagebox.showerror("שגיאה", f"שגיאה בבדיקת פריט שם לפני שמירת תמונה: {e}")
            return
        try:
            x = self.window.winfo_rootx() + self.canvas.winfo_x()
            y = self.window.winfo_rooty() + self.canvas.winfo_y()
            x1 = x + self.canvas.winfo_width()
            y1 = y + self.canvas.winfo_height()
            default_filename_img = f"מפה נומרולוגית - {self.full_name} - {self.entry_date.get()}.png"
            default_filename_img = "".join(
                c if c.isalnum() or c in (' ', '-', '_', '.', '(', ')', '[', ']') else '_' for c in
                default_filename_img)
            image_path = filedialog.asksaveasfilename(initialfile=default_filename_img, defaultextension=".png",
                                                      filetypes=[("PNG files", "*.png"), ("All Files", "*.*")])
            if image_path:
                ImageGrab.grab().crop((x, y, x1, y1)).save(image_path)
                messagebox.showinfo("הצלחה", f"התמונה נשמרה בהצלחה במיקום:\n{image_path}")
            else:
                messagebox.showinfo("בוטל", "שמירת התמונה בוטלה.")
        except Exception as e:
            error_msg = f"שגיאה במהלך שמירת התמונה:\n{e}\n{traceback.format_exc()}"
            print(f"Error saving image: {error_msg}")
            messagebox.showerror("שגיאת שמירת תמונה", error_msg)

# if __name__ == "__main__":
#     app = FullDates()
#     app.get_date()
